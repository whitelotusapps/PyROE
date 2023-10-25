import os
import sys
from datetime import datetime

import PyPDF2
from docx import Document
from docx.enum.style import WD_STYLE_TYPE  # Import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH  # Import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from PyQt5.QtGui import QFont
from PyQt5.QtWidgets import (QAction, QApplication, QComboBox, QFileDialog,
                             QLabel, QLineEdit, QListWidget, QMainWindow,
                             QPushButton, QTabWidget, QVBoxLayout,
                             QWidget)
class PDFGeneratorApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        menubar = self.menuBar()
        fileMenu = menubar.addMenu('File')

        generateAction = QAction('Generate ROE Documents', self)
        generateAction.triggered.connect(self.generate_pdf)
        fileMenu.addAction(generateAction)

        self.central_widget = QWidget(self)
        self.setCentralWidget(self.central_widget)
        layout = QVBoxLayout(self.central_widget)
        
        #########################################################################################################
        # SET GLOBAL FONT SIZE FOR APP UI
        #########################################################################################################
        font = QFont()
        font.setPointSize(14)  # Adjust the font size as needed
        self.setFont(font)
        #########################################################################################################

        #########################################################################################################
        #                      START OF CSS COLOR STYLING FOR TABS AND APP
        #########################################################################################################
        self.setStyleSheet('''
            QLabel {
                color: #333; /* Dark gray for label text */
                font-weight: bold;
                font-size: 14pt;
            }

            QLineEdit, QComboBox, QListWidget {
                border: 1px solid #333; /* Dark gray border for fields */
                font-size: 14pt;
            }
            QTabBar::tab {
                background-color: #f5f5f5; /* Light gray for tab background */
                font-size: 14pt; /* Set the font size for tab text */
                min-width: 220px;
            }
            QTabBar::tab:selected {
                background-color: #1e90ff; /* Blue for the selected tab */
                color: #fff;
            }
        ''')
        #########################################################################################################
        #                      END OF CSS COLOR STYLING FOR TABS AND APP
        #########################################################################################################
        
        tab_widget = QTabWidget()

        #########################################################################################################
        #                      START OF CREATING TABS / TAB LAYOUT
        #########################################################################################################

        #########################################################################################################
        # Tab 1: Personal Information
        #########################################################################################################
        tab1 = QWidget()
        tab1.layout = QVBoxLayout()
        self.first_given_name_label = QLabel('Enter your given first name:')
        self.first_given_name_text = QLineEdit()

        self.middle_given_name_label = QLabel('Enter your given middle name:')
        self.middle_given_name_text = QLineEdit()

        self.family_name_label = QLabel('Enter your family name (last name):')
        self.family_name_text = QLineEdit()

        self.street_address_label = QLabel('Your street address, spelled out. (Rd = Road, St = Street, etc):')
        self.street_address_text = QLineEdit()

        self.city_label = QLabel('Your address city:')
        self.city_text = QLineEdit()

        self.mailing_state_label = QLabel('Your address state:')
        self.mailing_state_combo = QComboBox()
        self.mailing_state_combo.addItems([
            "Alabama", "Alaska", "Arizona", "Arkansas", "California",
            "Colorado", "Connecticut", "Delaware", "Florida", "Georgia",
            "Hawaii", "Idaho", "Illinois", "Indiana", "Iowa", "Kansas",
            "Kentucky", "Louisiana", "Maine", "Maryland", "Massachusetts",
            "Michigan", "Minnesota", "Mississippi", "Missouri", "Montana",
            "Nebraska", "Nevada", "New Hampshire", "New Jersey", "New Mexico",
            "New York", "North Carolina", "North Dakota", "Ohio", "Oklahoma",
            "Oregon", "Pennsylvania", "Rhode Island", "South Carolina", "South Dakota",
            "Tennessee", "Texas", "Utah", "Vermont", "Virginia", "Washington",
            "West Virginia", "Wisconsin", "Wyoming"
        ])

        self.zip_label = QLabel('Your address zip code (without the [] backets):')
        self.zip_text = QLineEdit()

        self.social_security_number_label = QLabel("Social Security Number (xxx-xx-xxxx)")
        self.social_security_number_text = QLineEdit()

        self.man_or_woman_label = QLabel('Your gender:')
        self.man_or_woman = QComboBox()
        self.man_or_woman.addItems(['Man', 'Woman'])

        self.republic_of_birth_label = QLabel('Which state where you born in:')
        self.republic_of_birth = QComboBox()
        self.republic_of_birth.addItems([
            "Alabama", "Alaska", "Arizona", "Arkansas", "California",
            "Colorado", "Connecticut", "Delaware", "Florida", "Georgia",
            "Hawaii", "Idaho", "Illinois", "Indiana", "Iowa", "Kansas",
            "Kentucky", "Louisiana", "Maine", "Maryland", "Massachusetts",
            "Michigan", "Minnesota", "Mississippi", "Missouri", "Montana",
            "Nebraska", "Nevada", "New Hampshire", "New Jersey", "New Mexico",
            "New York", "North Carolina", "North Dakota", "Ohio", "Oklahoma",
            "Oregon", "Pennsylvania", "Rhode Island", "South Carolina", "South Dakota",
            "Tennessee", "Texas", "Utah", "Vermont", "Virginia", "Washington",
            "West Virginia", "Wisconsin", "Wyoming"
        ])        

        # ADD WIDGETS TO TAB LAYOUT
        tab1.layout.addWidget(self.first_given_name_label)
        tab1.layout.addWidget(self.first_given_name_text)
        tab1.layout.addWidget(self.middle_given_name_label)
        tab1.layout.addWidget(self.middle_given_name_text)
        tab1.layout.addWidget(self.family_name_label)
        tab1.layout.addWidget(self.family_name_text)
        tab1.layout.addWidget(self.street_address_label)
        tab1.layout.addWidget(self.street_address_text)
        tab1.layout.addWidget(self.city_label)
        tab1.layout.addWidget(self.city_text)
        tab1.layout.addWidget(self.mailing_state_label)
        tab1.layout.addWidget(self.mailing_state_combo)
        tab1.layout.addWidget(self.zip_label)
        tab1.layout.addWidget(self.zip_text)
        tab1.layout.addWidget(self.social_security_number_label)
        tab1.layout.addWidget(self.social_security_number_text)
        tab1.layout.addWidget(self.man_or_woman_label)
        tab1.layout.addWidget(self.man_or_woman)
        tab1.layout.addWidget(self.republic_of_birth_label)
        tab1.layout.addWidget(self.republic_of_birth)

        tab1.setLayout(tab1.layout)
        tab_widget.addTab(tab1, "Personal Information")

        #########################################################################################################
        # Tab 2: States of Sojourn
        #########################################################################################################
        tab2 = QWidget()
        tab2.layout = QVBoxLayout()
        self.states_label = QLabel('Which state(s) do you Sojourn in:')
        self.states_list = QListWidget()
        self.states_list.setSelectionMode(QListWidget.MultiSelection)
        self.states_list.addItems([
            "Alabama", "Alaska", "Arizona", "Arkansas", "California",
            "Colorado", "Connecticut", "Delaware", "Florida", "Georgia",
            "Hawaii", "Idaho", "Illinois", "Indiana", "Iowa", "Kansas",
            "Kentucky", "Louisiana", "Maine", "Maryland", "Massachusetts",
            "Michigan", "Minnesota", "Mississippi", "Missouri", "Montana",
            "Nebraska", "Nevada", "New Hampshire", "New Jersey", "New Mexico",
            "New York", "North Carolina", "North Dakota", "Ohio", "Oklahoma",
            "Oregon", "Pennsylvania", "Rhode Island", "South Carolina", "South Dakota",
            "Tennessee", "Texas", "Utah", "Vermont", "Virginia", "Washington",
            "West Virginia", "Wisconsin", "Wyoming"
        ])
        
        
        tab2.layout.addWidget(self.states_label)
        tab2.layout.addWidget(self.states_list)
        tab2.setLayout(tab2.layout)
        tab_widget.addTab(tab2, "States of Sojourn")

        #########################################################################################################
        # Tab 3: IRS Information
        #########################################################################################################
        tab3 = QWidget()
        tab3.layout = QVBoxLayout()
        self.irs_commissioner_label = QLabel('Current IRS Commissioner Name (Example: Daniel Werfel):')
        self.irs_commissioner_text = QLineEdit()
        self.local_irs_service_center_street_address_label = QLabel('Your local IRS service center street address:')
        self.local_irs_service_center_street_address_text = QLineEdit()
        self.local_irs_service_center_city_label = QLabel('Your local IRS service center city:')
        self.local_irs_service_center_city_text = QLineEdit()
        self.local_irs_service_center_state_label = QLabel('Your local IRS service center state:')
        self.local_irs_service_center_state_combo = QComboBox()
        self.local_irs_service_center_state_combo.addItems([
            "Alabama", "Alaska", "Arizona", "Arkansas", "California",
            "Colorado", "Connecticut", "Delaware", "Florida", "Georgia",
            "Hawaii", "Idaho", "Illinois", "Indiana", "Iowa", "Kansas",
            "Kentucky", "Louisiana", "Maine", "Maryland", "Massachusetts",
            "Michigan", "Minnesota", "Mississippi", "Missouri", "Montana",
            "Nebraska", "Nevada", "New Hampshire", "New Jersey", "New Mexico",
            "New York", "North Carolina", "North Dakota", "Ohio", "Oklahoma",
            "Oregon", "Pennsylvania", "Rhode Island", "South Carolina", "South Dakota",
            "Tennessee", "Texas", "Utah", "Vermont", "Virginia", "Washington",
            "West Virginia", "Wisconsin", "Wyoming"
        ])
        self.local_irs_service_center_zip_label = QLabel('Your local IRS service center zip code:')
        self.local_irs_service_center_zip_text = QLineEdit()
        
        tab3.layout.addWidget(self.irs_commissioner_label)
        tab3.layout.addWidget(self.irs_commissioner_text)
        tab3.layout.addWidget(self.local_irs_service_center_street_address_label)
        tab3.layout.addWidget(self.local_irs_service_center_street_address_text)
        tab3.layout.addWidget(self.local_irs_service_center_city_label)
        tab3.layout.addWidget(self.local_irs_service_center_city_text)
        tab3.layout.addWidget(self.local_irs_service_center_state_label)
        tab3.layout.addWidget(self.local_irs_service_center_state_combo)
        tab3.layout.addWidget(self.local_irs_service_center_zip_label)
        tab3.layout.addWidget(self.local_irs_service_center_zip_text)
        tab3.setLayout(tab3.layout)
        tab_widget.addTab(tab3, "IRS Information")

        #########################################################################################################
        # Tab 4: Notary Information
        #########################################################################################################
        tab4 = QWidget()
        tab4.layout = QVBoxLayout()

        self.notary_state_label = QLabel('State where affidavit is notarized:')
        self.notary_state_combo = QComboBox()
        self.notary_state_combo.addItems([
            "Alabama", "Alaska", "Arizona", "Arkansas", "California",
            "Colorado", "Connecticut", "Delaware", "Florida", "Georgia",
            "Hawaii", "Idaho", "Illinois", "Indiana", "Iowa", "Kansas",
            "Kentucky", "Louisiana", "Maine", "Maryland", "Massachusetts",
            "Michigan", "Minnesota", "Mississippi", "Missouri", "Montana",
            "Nebraska", "Nevada", "New Hampshire", "New Jersey", "New Mexico",
            "New York", "North Carolina", "North Dakota", "Ohio", "Oklahoma",
            "Oregon", "Pennsylvania", "Rhode Island", "South Carolina", "South Dakota",
            "Tennessee", "Texas", "Utah", "Vermont", "Virginia", "Washington",
            "West Virginia", "Wisconsin", "Wyoming"
        ])

        self.notary_county_label = QLabel('County where affidavit is notarized:')
        self.notary_county_text = QLineEdit()

        tab4.layout.addWidget(self.notary_state_label)
        tab4.layout.addWidget(self.notary_state_combo)
        tab4.layout.addWidget(self.notary_county_label)
        tab4.layout.addWidget(self.notary_county_text)
        tab4.setLayout(tab4.layout)
        tab_widget.addTab(tab4, "Notary Information")


        #########################################################################################################
        # Tab 5: W-8BEN PDF Information
        #########################################################################################################
        tab5 = QWidget()
        tab5.layout = QVBoxLayout()

        self.include_ohio_state_assembly_edits_label = QLabel('Would you like to adjust your W-8BEN to include Ohio State Assembly Edits?')
        self.include_ohio_state_assembly_edits_combo = QComboBox()
        self.include_ohio_state_assembly_edits_combo.addItems(["", "Yes", "No"])

        self.customize_ohio_state_assembly_text_label = QLabel('Would you like to customize any of the Ohio State Assembly Edit text?')
        self.customize_ohio_state_assembly_text_combo = QComboBox()
        self.customize_ohio_state_assembly_text_combo.addItems(["", "Yes", "No"])

        self.custom_ohio_state_assembly_text1_label = QLabel("Part 1, Identification of \"Non-Resident Alien\"")
        self.custom_ohio_state_assembly_text1_text = QLineEdit()

        self.custom_ohio_state_assembly_text2_label = QLabel("Part 1, Name of individual who is the \"Non-Resident Alien\"")
        self.custom_ohio_state_assembly_text2_text = QLineEdit()

        self.custom_ohio_state_assembly_text3_label = QLabel("Permanent \"domicile\" address")
        self.custom_ohio_state_assembly_text3_text = QLineEdit()

        self.custom_ohio_state_assembly_text4_label = QLabel("Margin text, \"Speical Private and Priority\". You will need to enter new lines (carriage returns) as \"\\n\"")
        self.custom_ohio_state_assembly_text4_text = QLineEdit()

        self.custom_ohio_state_assembly_text5_label = QLabel("Part 2, Field 9: \"non-resident alien owner\"")
        self.custom_ohio_state_assembly_text5_text = QLineEdit()

        self.custom_ohio_state_assembly_text6_label = QLabel("Part 3, \"non-resident alient owner\"")
        self.custom_ohio_state_assembly_text6_text = QLineEdit()

        self.custom_ohio_state_assembly_text7_label = QLabel("Signature line \"© All Right Reserved 28 U.S.C. § 1746(1)\"")
        self.custom_ohio_state_assembly_text7_text = QLineEdit()

        self.custom_ohio_state_assembly_text8_label = QLabel("Below signature line \"non-resident alient owner\"")
        self.custom_ohio_state_assembly_text8_text = QLineEdit()

        self.custom_ohio_state_assembly_text9_label = QLabel("Below signature line (or individual authorized to sign for \"non-resident alient owner\")")
        self.custom_ohio_state_assembly_text9_text = QLineEdit()

        self.custom_ohio_state_assembly_text10_label = QLabel("Print name of signer line")
        self.custom_ohio_state_assembly_text10_text = QLineEdit()


        # Setting default text for custom Ohio State Assembly Edit text fields
        self.custom_ohio_state_assembly_text1_text.setText("Non-Resident Alien")
        self.custom_ohio_state_assembly_text2_text.setText("Non-Resident Alien")
        self.custom_ohio_state_assembly_text3_text.setText("  domicile")
        self.custom_ohio_state_assembly_text4_text.setText("Speical\n\nPrivate\n\nand\n\nPriority")
        self.custom_ohio_state_assembly_text5_text.setText("non-resident alien owner")
        self.custom_ohio_state_assembly_text6_text.setText("non-resident alien owner")
        self.custom_ohio_state_assembly_text7_text.setText("© All Right Reserved 28 U.S.C. § 1746(1)")
        self.custom_ohio_state_assembly_text8_text.setText("non-resident alient owner")
        self.custom_ohio_state_assembly_text9_text.setText("non-resident alient owner")
        self.custom_ohio_state_assembly_text10_text.setText("In Exclusive Equity")

        self.date_of_birth_label = QLabel('What is your date of birth? (MM/DD/YYYY)')
        self.date_of_birth_text = QLineEdit()

        tab5.layout.addWidget(self.include_ohio_state_assembly_edits_label)
        tab5.layout.addWidget(self.include_ohio_state_assembly_edits_combo)
        tab5.layout.addWidget(self.customize_ohio_state_assembly_text_label)
        tab5.layout.addWidget(self.customize_ohio_state_assembly_text_combo)

        tab5.layout.addWidget(self.custom_ohio_state_assembly_text1_label)
        tab5.layout.addWidget(self.custom_ohio_state_assembly_text1_text)

        tab5.layout.addWidget(self.custom_ohio_state_assembly_text2_label)
        tab5.layout.addWidget(self.custom_ohio_state_assembly_text2_text)

        tab5.layout.addWidget(self.custom_ohio_state_assembly_text3_label)
        tab5.layout.addWidget(self.custom_ohio_state_assembly_text3_text)

        tab5.layout.addWidget(self.custom_ohio_state_assembly_text4_label)
        tab5.layout.addWidget(self.custom_ohio_state_assembly_text4_text)

        tab5.layout.addWidget(self.custom_ohio_state_assembly_text5_label)
        tab5.layout.addWidget(self.custom_ohio_state_assembly_text5_text)

        tab5.layout.addWidget(self.custom_ohio_state_assembly_text6_label)
        tab5.layout.addWidget(self.custom_ohio_state_assembly_text6_text)

        tab5.layout.addWidget(self.custom_ohio_state_assembly_text7_label)
        tab5.layout.addWidget(self.custom_ohio_state_assembly_text7_text)

        tab5.layout.addWidget(self.custom_ohio_state_assembly_text8_label)
        tab5.layout.addWidget(self.custom_ohio_state_assembly_text8_text)

        tab5.layout.addWidget(self.custom_ohio_state_assembly_text9_label)
        tab5.layout.addWidget(self.custom_ohio_state_assembly_text9_text)

        tab5.layout.addWidget(self.custom_ohio_state_assembly_text10_label)
        tab5.layout.addWidget(self.custom_ohio_state_assembly_text10_text)

        tab5.layout.addWidget(self.date_of_birth_label)
        tab5.layout.addWidget(self.date_of_birth_text)

        # Connect the first question's combobox to a slot
        self.include_ohio_state_assembly_edits_combo.currentIndexChanged.connect(self.toggleCustomizeText)

        # Initially hide the second question
        self.customize_ohio_state_assembly_text_label.hide()
        self.customize_ohio_state_assembly_text_combo.hide()

        self.custom_ohio_state_assembly_text1_label.hide()
        self.custom_ohio_state_assembly_text1_text.hide()

        self.custom_ohio_state_assembly_text2_label.hide()
        self.custom_ohio_state_assembly_text2_text.hide()

        self.custom_ohio_state_assembly_text3_label.hide()
        self.custom_ohio_state_assembly_text3_text.hide()

        self.custom_ohio_state_assembly_text4_label.hide()
        self.custom_ohio_state_assembly_text4_text.hide()

        self.custom_ohio_state_assembly_text5_label.hide()
        self.custom_ohio_state_assembly_text5_text.hide()

        self.custom_ohio_state_assembly_text6_label.hide()
        self.custom_ohio_state_assembly_text6_text.hide()

        self.custom_ohio_state_assembly_text7_label.hide()
        self.custom_ohio_state_assembly_text7_text.hide()

        self.custom_ohio_state_assembly_text8_label.hide()
        self.custom_ohio_state_assembly_text8_text.hide()

        self.custom_ohio_state_assembly_text9_label.hide()
        self.custom_ohio_state_assembly_text9_text.hide()

        self.custom_ohio_state_assembly_text10_label.hide()
        self.custom_ohio_state_assembly_text10_text.hide()


        # Connect the first question's combobox to a slot
        self.customize_ohio_state_assembly_text_combo.currentIndexChanged.connect(self.toggleCustomizeTextOptions)

        tab5.setLayout(tab5.layout)
        tab_widget.addTab(tab5, "W-8BEN PDF Info")

        #########################################################################################################
        #                       END OF CREATING TABS / TAB LAYOUT
        #########################################################################################################

        layout.addWidget(tab_widget)

        self.generate_button = QPushButton('Generate ROE Documents')
        layout.addWidget(self.generate_button)

        self.generate_button.clicked.connect(self.generate_pdf)

        #########################################################################################################
        #                       START OF SCREEN SIZE DETECTION (to center the app on the screen)
        #########################################################################################################
        self.setWindowTitle('PyROE v0.2')
        self.showMaximized()
        #########################################################################################################
        #                       END OF SCREEN SIZE DETECTION (to center the app on the screen)
        #########################################################################################################

    def toggleCustomizeText(self, index):
        if index == 1:  # If "Yes" is selected in the first question
            self.customize_ohio_state_assembly_text_label.show()
            self.customize_ohio_state_assembly_text_combo.show()
        else:
            self.customize_ohio_state_assembly_text_label.hide()
            self.customize_ohio_state_assembly_text_combo.hide()
            self.customize_ohio_state_assembly_text_label.hide()
            self.customize_ohio_state_assembly_text_combo.hide()

            self.custom_ohio_state_assembly_text1_label.hide()
            self.custom_ohio_state_assembly_text1_text.hide()

            self.custom_ohio_state_assembly_text2_label.hide()
            self.custom_ohio_state_assembly_text2_text.hide()

            self.custom_ohio_state_assembly_text3_label.hide()
            self.custom_ohio_state_assembly_text3_text.hide()

            self.custom_ohio_state_assembly_text4_label.hide()
            self.custom_ohio_state_assembly_text4_text.hide()

            self.custom_ohio_state_assembly_text5_label.hide()
            self.custom_ohio_state_assembly_text5_text.hide()

            self.custom_ohio_state_assembly_text6_label.hide()
            self.custom_ohio_state_assembly_text6_text.hide()

            self.custom_ohio_state_assembly_text7_label.hide()
            self.custom_ohio_state_assembly_text7_text.hide()

            self.custom_ohio_state_assembly_text8_label.hide()
            self.custom_ohio_state_assembly_text8_text.hide()

            self.custom_ohio_state_assembly_text9_label.hide()
            self.custom_ohio_state_assembly_text9_text.hide()

            self.custom_ohio_state_assembly_text10_label.hide()
            self.custom_ohio_state_assembly_text10_text.hide()

    def toggleCustomizeTextOptions(self, index):
        if index == 1:  # If "Yes" is selected in the first question
            self.customize_ohio_state_assembly_text_label.show()
            self.customize_ohio_state_assembly_text_combo.show()
            # Show the additional labels for customizing Ohio State Assembly Edit text
            self.custom_ohio_state_assembly_text1_label.show()
            self.custom_ohio_state_assembly_text1_text.show()

            self.custom_ohio_state_assembly_text2_label.show()
            self.custom_ohio_state_assembly_text2_text.show()

            self.custom_ohio_state_assembly_text3_label.show()
            self.custom_ohio_state_assembly_text3_text.show()

            self.custom_ohio_state_assembly_text4_label.show()
            self.custom_ohio_state_assembly_text4_text.show()

            self.custom_ohio_state_assembly_text5_label.show()
            self.custom_ohio_state_assembly_text5_text.show()

            self.custom_ohio_state_assembly_text6_label.show()
            self.custom_ohio_state_assembly_text6_text.show()

            self.custom_ohio_state_assembly_text7_label.show()
            self.custom_ohio_state_assembly_text7_text.show()

            self.custom_ohio_state_assembly_text8_label.show()
            self.custom_ohio_state_assembly_text8_text.show()

            self.custom_ohio_state_assembly_text9_label.show()
            self.custom_ohio_state_assembly_text9_text.show()

            self.custom_ohio_state_assembly_text10_label.show()
            self.custom_ohio_state_assembly_text10_text.show()
        else:
            self.customize_ohio_state_assembly_text_label.hide()
            self.customize_ohio_state_assembly_text_combo.hide()
            # Hide the additional labels for customizing Ohio State Assembly Edit text
            self.custom_ohio_state_assembly_text1_label.hide()
            self.custom_ohio_state_assembly_text1_text.hide()

            self.custom_ohio_state_assembly_text2_label.hide()
            self.custom_ohio_state_assembly_text2_text.hide()

            self.custom_ohio_state_assembly_text3_label.hide()
            self.custom_ohio_state_assembly_text3_text.hide()

            self.custom_ohio_state_assembly_text4_label.hide()
            self.custom_ohio_state_assembly_text4_text.hide()

            self.custom_ohio_state_assembly_text5_label.hide()
            self.custom_ohio_state_assembly_text5_text.hide()

            self.custom_ohio_state_assembly_text6_label.hide()
            self.custom_ohio_state_assembly_text6_text.hide()

            self.custom_ohio_state_assembly_text7_label.hide()
            self.custom_ohio_state_assembly_text7_text.hide()

            self.custom_ohio_state_assembly_text8_label.hide()
            self.custom_ohio_state_assembly_text8_text.hide()

            self.custom_ohio_state_assembly_text9_label.hide()
            self.custom_ohio_state_assembly_text9_text.hide()

            self.custom_ohio_state_assembly_text10_label.hide()
            self.custom_ohio_state_assembly_text10_text.hide()

    def generate_pdf(self):

        #########################################################################################################
        # THE generate_state_dicts FUNCTION IS FOR TAKING THE STATE NAMES SELECTED FROM THE VARIOUS DROP DOWN MENUS
        # AND CREATE A DICTIONARY THAT IS USED BY SUBSEQUENT FUNCTIONS TO ACCESS REPUBLIC NAMES, STATE NAMES, AND STATE ABBREVIATIONS
        #########################################################################################################
        def generate_state_dicts(states):
            state_abbreviations = {
                "Alabama": ("Alabama Republic", "Alabama", "AL", "Alabama National (Alabaman)"),
                "Alaska": ("Alaska Republic", "Alaska", "AK", "Alaskan National (Alaskan)"),
                "Arizona": ("Arizona Republic", "Arizona", "AZ", "Arizona National (Arizonan)"),
                "Arkansas": ("Arkansas Republic", "Arkansas", "AR", "Arkansas National (Arkansan)"),
                "California": ("California Republic", "California", "CA", "California National (Californian)"),
                "Colorado": ("Colorado Republic", "Colorado", "CO", "Colorado National (Coloradan)"),
                "Connecticut": ("Connecticut Republic", "Connecticut", "CT", "Connecticut National (Connecticuter)"),
                "Delaware": ("Delaware Republic", "Delaware", "DE", "Delaware National (Delawarian)"),
                "Florida": ("Florida Republic", "Florida", "FL", "Florida National (Floridian)"),
                "Georgia": ("Georgia Republic", "Georgia", "GA", "Georgia National (Georgian)"),
                "Hawaii": ("Hawaii Republic", "Hawaii", "HI", "Hawaii National (Hawaiian)"),
                "Idaho": ("Idaho Republic", "Idaho", "ID", "Idaho National (Idahoan)"),
                "Illinois": ("Illinois Republic", "Illinois", "IL", "Illinois National (Illinoisan)"),
                "Indiana": ("Indiana Republic", "Indiana", "IN", "Indiana National (Hoosier)"),
                "Iowa": ("Iowa Republic", "Iowa", "IA", "Iowa National (Iowan)"),
                "Kansas": ("Kansas Republic", "Kansas", "KS", "Kansas National (Kansan)"),
                "Kentucky": ("Kentucky Republic", "Kentucky", "KY", "Kentucky National (Kentuckian)"),
                "Louisiana": ("Louisiana Republic", "Louisiana", "LA", "Louisiana National (Louisianan)"),
                "Maine": ("Maine Republic", "Maine", "ME", "Maine National (Mainer)"),
                "Maryland": ("Maryland Republic", "Maryland", "MD", "Maryland National (Marylander)"),
                "Massachusetts": ("Massachusetts Republic", "Massachusetts", "MA", "Massachusetts National (Bay Stater)"),
                "Michigan": ("Michigan Republic", "Michigan", "MI", "Michigan National (Michigander)"),
                "Minnesota": ("Minnesota Republic", "Minnesota", "MN", "Minnesota National (Minnesotan)"),
                "Mississippi": ("Mississippi Republic", "Mississippi", "MS", "Mississippi National (Mississippian)"),
                "Missouri": ("Missouri Republic", "Missouri", "MO", "Missouri National (Missourian)"),
                "Montana": ("Montana Republic", "Montana", "MT", "Montana National (Montanan)"),
                "Nebraska": ("Nebraska Republic", "Nebraska", "NE", "Nebraska National (Nebraskan)"),
                "Nevada": ("Nevada Republic", "Nevada", "NV", "Nevada National (Nevadan)"),
                "New Hampshire": ("New Hampshire Republic", "New Hampshire", "NH", "New Hampshire National (New Hampshirite)"),
                "New Jersey": ("New Jersey Republic", "New Jersey", "NJ", "New Jersey National (New Jerseyan)"),
                "New Mexico": ("New Mexico Republic", "New Mexico", "NM", "New Mexico National (New Mexican)"),
                "New York": ("New York Republic", "New York", "NY", "New York National (New Yorker)"),
                "North Carolina": ("North Carolina Republic", "North Carolina", "NC", "North Carolina National (North Carolinian)"),
                "North Dakota": ("North Dakota Republic", "North Dakota", "ND", "North Dakota National (North Dakotan)"),
                "Ohio": ("Ohio Republic", "Ohio", "OH", "Ohio National (Ohioan)"),
                "Oklahoma": ("Oklahoma Republic", "Oklahoma", "OK", "Oklahoma National (Oklahoman)"),
                "Oregon": ("Oregon Republic", "Oregon", "OR", "Oregon National (Oregonian)"),
                "Pennsylvania": ("Pennsylvania Republic", "Pennsylvania", "PA", "Pennsylvania National (Pennsylvanian)"),
                "Rhode Island": ("Rhode Island Republic", "Rhode Island", "RI", "Rhode Island National (Rhode Islander)"),
                "South Carolina": ("South Carolina Republic", "South Carolina", "SC", "South Carolina National (South Carolinian)"),
                "South Dakota": ("South Dakota Republic", "South Dakota", "SD", "South Dakota National (South Dakotan)"),
                "Tennessee": ("Tennessee Republic", "Tennessee", "TN", "Tennessee National (Tennessean)"),
                "Texas": ("Texas Republic", "Texas", "TX", "Texas National (Texan)"),
                "Utah": ("Utah Republic", "Utah", "UT", "Utah National (Utahn)"),
                "Vermont": ("Vermont Republic", "Vermont", "VT", "Vermont National (Vermonter)"),
                "Virginia": ("Virginia Republic", "Virginia", "VA", "Virginia National (Virginian)"),
                "Washington": ("Washington Republic", "Washington", "WA", "Washington National (Washingtonian)"),
                "West Virginia": ("West Virginia Republic", "West Virginia", "WV", "West Virginia National (West Virginian)"),
                "Wisconsin": ("Wisconsin Republic", "Wisconsin", "WI", "Wisconsin National (Wisconsinite)"),
                "Wyoming": ("Wyoming Republic", "Wyoming", "WY", "Wyoming National (Wyomingite)")
            }

            if isinstance(states, str):
                # If a single string is passed, convert it to a list
                states = [states]

            state_dicts = []
            for state in states:
                state_info = state_abbreviations.get(state, ("", "", "", ""))
                state_dict = {
                    "Republic Name": state_info[0],
                    "State Name": state_info[1],
                    "State Abbreviation": state_info[2],
                    "Country of Citizenship": state_info[3]
                }
                state_dicts.append(state_dict)
            return state_dicts
        
        
            #########################################################################################################
            #       END OF generate_state_dicts(states)
            #########################################################################################################

        #########################################################################################################
        # START OF DEFINE VARIABLES
        #########################################################################################################
        first_given_name = self.first_given_name_text.text()
        middle_given_name = self.middle_given_name_text.text()
        family_name = self.family_name_text.text()
        man_or_woman = self.man_or_woman.currentText()
        selected_states = [item.text() for item in self.states_list.selectedItems()]
        street_address = self.street_address_text.text()
        city = self.city_text.text()
        zip_code = self.zip_text.text()
        mailing_state = self.mailing_state_combo.currentText()
        social_security_number = self.social_security_number_text.text()
        irs_commissioner = self.irs_commissioner_text.text()
        local_irs_service_center_street_address = self.local_irs_service_center_street_address_text.text()
        local_irs_service_center_city = self.local_irs_service_center_city_text.text()
        local_irs_service_center_state = self.local_irs_service_center_state_combo.currentText()
        local_irs_service_center_zip = self.local_irs_service_center_zip_text.text()
        republic_of_birth = self.republic_of_birth.currentText()
        notary_state = self.notary_state_combo.currentText()
        notary_county = self.notary_county_text.text()
        sojourn_states_list = generate_state_dicts(selected_states)
        mailing_address_state_list = generate_state_dicts(mailing_state)
        republic_of_birth_list = generate_state_dicts(republic_of_birth)
        notary_state_list = generate_state_dicts(notary_state)
        local_irs_service_center_state_list = generate_state_dicts(local_irs_service_center_state)
        timestamp = datetime.now().strftime("%H_%M_%S")
        letter_of_intent_filename = f"01 - {first_given_name} {middle_given_name} {family_name} - Letter of Intent - {timestamp}.docx"
        affidavit_filename = f"02 - {first_given_name} {middle_given_name} {family_name} - Affidavit - {timestamp}.docx"
        supporting_evidence_filename = f"03 - {first_given_name} {middle_given_name} {family_name} - Supporting Evidence - {timestamp}.docx"
        w_8ben_pdf_filename = f"04 - {first_given_name} {middle_given_name} {family_name} - Supporting Evidence - {timestamp}.pdf"
        country_of_citizenship = republic_of_birth_list[0]['Country of Citizenship']
        date_of_birth = self.date_of_birth_text.text()
        include_ohio_state_edits = self.include_ohio_state_assembly_edits_combo.currentText()
        #########################################################################################################
        # END OF DEFINE VARIABLES
        #########################################################################################################

        def any_variable_empty(*variables):
            return any(not var for var in variables)

        if any_variable_empty(first_given_name, middle_given_name, family_name, man_or_woman, selected_states, street_address, city, zip_code, mailing_state, social_security_number, irs_commissioner, local_irs_service_center_street_address, local_irs_service_center_city, local_irs_service_center_state, local_irs_service_center_zip, republic_of_birth, notary_state, notary_county):
            return
        
        docx_folder_path = QFileDialog.getExistingDirectory(self, "Select Folder to Save Your Revocation of Election Files", os.path.expanduser("~"))
        
        if docx_folder_path:

            # LETTER OF INTENT:
            letter_of_intent_file_path = os.path.join(docx_folder_path, letter_of_intent_filename)
            letter_of_intent_document = Document()
            self.create_letter_of_intent(letter_of_intent_document, first_given_name, middle_given_name, family_name, man_or_woman, selected_states, street_address, city, zip_code, mailing_state, social_security_number, irs_commissioner, local_irs_service_center_street_address, local_irs_service_center_city, local_irs_service_center_state_list, local_irs_service_center_zip, sojourn_states_list, mailing_address_state_list, republic_of_birth_list, notary_state_list, notary_county)
            letter_of_intent_document.save(letter_of_intent_file_path)
            
            # AFFIDAVIT:
            affidavit_file_path = os.path.join(docx_folder_path, affidavit_filename)
            affidavit_document = Document()
            self.create_affidavit(affidavit_document, first_given_name, middle_given_name, family_name, man_or_woman, selected_states, street_address, city, zip_code, mailing_state, social_security_number, irs_commissioner, local_irs_service_center_street_address, local_irs_service_center_city, local_irs_service_center_state_list, local_irs_service_center_zip, republic_of_birth, sojourn_states_list, mailing_address_state_list, republic_of_birth_list, notary_state_list, notary_county)
            affidavit_document.save(affidavit_file_path)
            
            # SUPPORTING EVIDENCE:
            supporting_evidence_file_path = os.path.join(docx_folder_path, supporting_evidence_filename)
            supporting_evidence_document = Document()
            self.create_supporting_evidence(supporting_evidence_document, first_given_name, middle_given_name, family_name, man_or_woman, selected_states, street_address, city, zip_code, mailing_state, social_security_number, irs_commissioner, local_irs_service_center_street_address, local_irs_service_center_city, local_irs_service_center_state_list, local_irs_service_center_zip, sojourn_states_list, mailing_address_state_list, republic_of_birth_list, notary_state_list, notary_county)
            supporting_evidence_document.save(supporting_evidence_file_path)

            # W-8BEN PDF CREATION:
            w_8ben_pdf_file_path = os.path.join(docx_folder_path, w_8ben_pdf_filename)
            self.create_w_8ben_pdf(w_8ben_pdf_file_path, first_given_name, middle_given_name, family_name, man_or_woman, selected_states, street_address, city, zip_code, mailing_state, social_security_number, irs_commissioner, local_irs_service_center_street_address, local_irs_service_center_city, local_irs_service_center_state_list, local_irs_service_center_zip, sojourn_states_list, mailing_address_state_list, republic_of_birth_list, notary_state_list, notary_county, country_of_citizenship, date_of_birth, include_ohio_state_edits)

    #########################################################################################################
    #                       START OF create_letter_of_intent FUNCTION
    #########################################################################################################
    def create_letter_of_intent(self, docx_document, first_given_name, middle_given_name, family_name, man_or_woman, selected_states, street_address, city, zip_code, mailing_state, social_security_number, irs_commissioner, local_irs_service_center_street_address, local_irs_service_center_city, local_irs_service_center_state_list, local_irs_service_center_zip, sojourn_states_list, mailing_address_state_list, republic_of_birth_list, notary_state_list, notary_county):


        #########################################################################################################
        #                       START OF DEFINING VARIABLES
        #########################################################################################################
        # Extract the "State Name" from each dictionary
        soujourn_state_names_list = [state_dict["State Name"] for state_dict in sojourn_states_list]
        soujourn_state_names = ", ".join(soujourn_state_names_list)

        # Extract the "State Abbreviations" from each dictionary
        soujourn_state_abbreviations_list = [state_dict["State Abbreviation"] for state_dict in sojourn_states_list]
        soujourn_state_abbreviations = ", ".join(soujourn_state_abbreviations_list)

        # Extract the "State " from each dictionary
        soujourn_republic_name_list = [state_dict["Republic Name"] for state_dict in sojourn_states_list]
        soujourn_republic_names = ", ".join(soujourn_republic_name_list)

        mailing_address_republic_name = mailing_address_state_list[0]['Republic Name']
        mailing_address_state_name = mailing_address_state_list[0]['State Name']
        mailing_address_state_abbreviation = mailing_address_state_list[0]['State Abbreviation']

        republic_of_birth_name = republic_of_birth_list[0]['Republic Name']
        republic_of_birth_state = republic_of_birth_list[0]['State Name']
        republic_of_birth_state_abbreviation = republic_of_birth_list[0]['State Abbreviation']

        local_irs_service_center_republic_name = local_irs_service_center_state_list[0]['Republic Name']
        local_irs_service_center_state__name = local_irs_service_center_state_list[0]['State Name']
        local_irs_service_center_state_abbreviation = local_irs_service_center_state_list[0]['State Abbreviation']

        notary_republic_name = notary_state_list[0]['Republic Name']
        notary_state_name = notary_state_list[0]['State Name']
        notary_state_abbrivation = notary_state_list[0]['State Abbreviation']
        
        vessel_name = first_given_name + ' ' + middle_given_name + ' ' + family_name
        live_name = first_given_name + '-' + middle_given_name + ': ' + family_name

        # Determine the pronouns based on the 'are_you' variable
        if man_or_woman == 'Man':
            pronouns = ("he", "his", "him")
            gender = man_or_woman.lower()
        elif man_or_woman == 'Woman':
            pronouns = ("she", "her", "her")
            gender = man_or_woman.lower()


        # Create variables for titlecased and uppercase versions of the name
        titlecased_name = vessel_name.title()
        uppercase_name = vessel_name.upper()
        #########################################################################################################
        #                       END OF DEFINING VARIABLES
        #########################################################################################################


        #########################################################################################################
        #               START OF CREATING HEADER
        #########################################################################################################
        # Create a header for all pages
        section = docx_document.sections[0]
        header = section.header

        # Add the header text
        header_text = f"Revocation of Election for {titlecased_name}"
        header_paragraph = header.paragraphs[0]
        header_run = header_paragraph.add_run(header_text)
        
        # HEADER FONT SIZE
        header_run.font.size = Pt(12)
        header_run.bold = True
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_paragraph.space_after = Pt(12)  # Adjust the spacing as needed
        #########################################################################################################
        #               END OF CREATING HEADER
        #########################################################################################################

        #########################################################################################################
        #               START OF CREATING CUSTOM DOCUMENT STYLES
        #########################################################################################################
        # Create a custom style for normal document body text with a custom leading (line spacing)
        normal_style = docx_document.styles['Normal']
        normal_style.font.size = Pt(10)
        normal_style.paragraph_format.space_after = Pt(0)

        # Define a custom style with a left indent of 36 points
        indent_list_level_1 = docx_document.styles.add_style('IndentListLevel1', WD_STYLE_TYPE.PARAGRAPH)
        indent_list_level_1.font.size = Pt(10)
        indent_list_level_1.paragraph_format.space_after = Pt(0)
        indent_list_level_1.paragraph_format.left_indent = Pt(36)

        # Define a custom style with a left indent of 44 points
        indent_list_level_2 = docx_document.styles.add_style('IndentListLevel2', WD_STYLE_TYPE.PARAGRAPH)
        indent_list_level_2.font.size = Pt(10)
        indent_list_level_2.paragraph_format.space_after = Pt(0)
        indent_list_level_2.paragraph_format.left_indent = Pt(44)

        # Create a custom style for centered text
        centered_style = docx_document.styles.add_style('Centered', WD_STYLE_TYPE.PARAGRAPH)
        centered_style.font.size = Pt(10)
        centered_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        centered_style.paragraph_format.space_after = Pt(0)
        #########################################################################################################
        #               END OF CREATING CUSTOM DOCUMENT STYLES
        #########################################################################################################

        # Get the current date in the desired format
        current_date = datetime.now().strftime("%m/%d/%Y")

        #########################################################################################################
        #               START OF CREATING TABLE FOR LETTER OF INTENT
        #########################################################################################################
        table = docx_document.add_table(rows=6, cols=2)
        table.autofit = False
        table.allow_autofit = False
        table.columns[0].width = Pt(225)
        table.columns[1].width = Pt(225)

        # Remove all borders from the table
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        table.cell(0, 0).text = titlecased_name
        table.cell(0, 1).text = ""
        table.cell(1, 0).text = "c/o Non-domestic"
        table.cell(1, 1).text = f"Re: Social Security Number: {social_security_number}"
        table.cell(2, 0).text = street_address
        table.cell(2, 1).text = "assigned by a federal agency to"
        table.cell(3, 0).text = f"{city}, {mailing_state}"
        # Get the cell you want to modify
        cell = table.cell(3, 1)
        # Create a run for the normal text (not italicized)
        normal_run = cell.paragraphs[0].add_run(uppercase_name+ ", ")
        # Create a run for the italicized text
        italic_run = cell.paragraphs[0].add_run("nom-de-guerre")
        # Apply italic formatting to the italicized run
        italic_run.italic = True
        # Set the font size for both runs (adjust as needed)
        normal_run.font.size = Pt(10)
        italic_run.font.size = Pt(10)
        table.cell(4, 0).text = "USA [near " + zip_code + "]"
        table.cell(4, 1).text = "vessel, securitized negotiable instrument"
        table.autofit = False
        table.allow_autofit = False
        table.columns[0].width = Pt(225)
        table.columns[1].width = Pt(225)
        #########################################################################################################
        #               END OF CREATING TABLE FOR LETTER OF INTENT
        #########################################################################################################

        # Add a spacer (blank line) to create a new line
        docx_document.add_paragraph()

        # Add the current date
        docx_document.add_paragraph(f'Date: {current_date}', style=normal_style)

        # Add a spacer (blank line) to create a new line
        docx_document.add_paragraph()

        #########################################################################################################
        #               START OF IRS INFORMATION
        #########################################################################################################
        # Add IRS Commissioner information
        docx_document.add_paragraph(irs_commissioner, style=normal_style)
        docx_document.add_paragraph('IRS Commissioner', style=normal_style)
        docx_document.add_paragraph('Internal Revenue Service', style=normal_style)
        docx_document.add_paragraph('1111 Constitution Avenue NW', style=normal_style)
        docx_document.add_paragraph('Washington, DC 20224', style=normal_style)

        # Add a spacer (blank line) to create a new line
        docx_document.add_paragraph()

        # Add local IRS service center information
        docx_document.add_paragraph('DIRECTOR', style=normal_style)
        docx_document.add_paragraph('IRS Service Center', style=normal_style)
        docx_document.add_paragraph(local_irs_service_center_street_address, style=normal_style)
        docx_document.add_paragraph(local_irs_service_center_city + ', ' + local_irs_service_center_state_abbreviation + ' ' + local_irs_service_center_zip, style=normal_style)
        #########################################################################################################
        #               END OF IRS INFORMATION
        #########################################################################################################

        docx_document.add_paragraph()

        #########################################################################################################
        #               START OF LETTER OF INTENT
        #########################################################################################################
        docx_document.add_paragraph('REVOCATION OF ELECTION', style=centered_style)
        docx_document.add_paragraph('as established by the Congress of the United States', style=centered_style)

        docx_document.add_paragraph()

        docx_document.add_paragraph(f"Dear {irs_commissioner} and DIRECTOR,", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"Greetings! {titlecased_name}, a living {gender} and Affiant, is contacting you herein to give you Notice of {pronouns[1]} Revocation of Election, revoking {pronouns[1]} status as a federal U.S. citizen “taxpayer.” Affiant will no longer be “volunteering” or “electing” to be treated as a federal taxpayer or U.S. citizen or as a “surety” for any government created entity connected to Form 1040 federal income tax “contributions.”", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"If you are not the appropriate IRS Officer with the duty to expedite this matter, please notify them immediately in writing and please forward this ROE to the IRS officer who is responsible for making the required changes to {pronouns[1]} taxable status.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"The enclosed documents, collectively referred to herein as {pronouns[1]} Revocation of Election (hereinafter ROE), consist of the following three (3) documents:", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"Notice of a ROE authorizing you or your delegates to “revoke” and “terminate” the IRS’s previously presumed “election” of Affiant to be taxed as a federal “U.S. citizen” and then to change {pronouns[1]} present tax status from a “taxpayer” to that of a “non-taxpayer” or a similar “non-taxable” term;", style='List Number 2')
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"AFFIDAVIT in support of and validating Affiant’s right to said tax status change and instructions for your use in executing the tax status change required to be made by laws passed by Congress.", style='List Number 2')
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"Supplemental information related to Affiant’s American National status and the laws passed by Congress affecting Affiant’s “non-taxable” status.", style='List Number 2')
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"The purpose of the three documents enclosed is to provide you Actual Notice, instructions, and Affiant’s consent to change Affiant’s tax status by terminating the IRS’s “presumed” sub silentio voluntary election to be treated as a “taxpayer” and to have you change Affiant’s tax status within all relevant IRS databases.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"Affiant’s AFFIDAVIT is enclosed and incorporated here for the purpose of explaining and verifying Affiant’s tax status jurisdiction that lawfully qualifies {pronouns[1]} to have {pronouns[1]} tax status changed to that of a “non-taxpayer” or any other similar non-taxable designation.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"Affiant believes in paying all taxes lawfully owed to support {pronouns[1]} country and {pronouns[0]} presently pays various types of sales taxes, use taxes, and State taxes to perform {pronouns[1]} civic duty to society. Affiant does not protest against any taxes lawfully owed.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"From years of research pertaining to the IRC, Title 26 U.S.C. and various U.S. Supreme Court and Appellate Court tax case rulings, it has come to Affiant’s attention that {pronouns[0]} has the qualifications to legally change {pronouns[1]} tax status to that of someone who is neither “subject to” nor “liable for” federal income taxes imposed by Subtitle A of the IRC.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"According to the IRC, higher Court case rulings, congressionally passed laws, and Affiant’s foreign jurisdiction to the IRS’s jurisdiction, Affiant has no legal obligation to file a Form 1040 as an American National {gender} not within (without) D.C.’s foreign jurisdiction.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"In particular, the title name of Form 1040 is: U.S. Individual Income Tax Return. The definitions at IRC 7701(a)(1) and 7701(a)(30), and the implementing Regulation at 26 CFR 1.1-1 make it quite clear that Form 1040 is intended to be executed only by U.S. citizens and resident aliens domiciled in the “United States.”", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"Affiant is neither a federal “U.S. citizen” nor a “resident alien.” Because IRC 6065 requires Form 1040 to be executed under penalties of perjury, it would be an act of felony perjury for Affiant to declare himself a “U.S. citizen” or “resident alien” by signing a Form 1040 when Affiant belongs to neither class of people. “U.S.” means “United States” which is legally defined in 26 U.S.C. to mean the District of Columbia (D.C.) and Affiant is not a D.C. resident or an “Individual” and therefore, is not legally qualified to file a Form 1040 and doing so would be an act of perjury. Affiant cannot file a Form 1040 as that would be an illegal act of “impersonating” a federal “U.S. citizen” or “U.S. person” in a foreign jurisdiction to Affiant's.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"Affiant has extensively reviewed the IRC, including many consultations with professional accountants, tax attorneys and IRS agents, and Affiant has unquestionably realized the truth that {pronouns[0]} has the inalienable right to lawfully discontinue volunteering to file Form 1040, based on {pronouns[1]} “non-federal,” “non U.S. citizen” and “not-qualified” status and it would be an act of felony perjury to sign and file a Form 1040 (see 18 USC 911&912). Form 1040 is not meant to be filed or legally required to be sent to the IRS by living men and woman American Nationals or state Citizens of the union as they are not “federal” U.S. citizens.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"No attorney, CPA tax accountant or IRS agent has ever been able to show Affiant where in the IRC or 26 U.S.C it says {pronouns[0]} is personally “liable for” the “federal” U.S. Individual income tax (unless one volunteers and “elects” to be treated as a federal  taxpayer even when the law says no tax liability is due).", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"Affiant discloses herein information related to the Chapters and sections of the IRC, 26 U.S.C., the original Constitution of 1787 and tax case rulings of the U.S. Supreme Court and other American Appellate Courts that substantiate Affiant’s findings and declarations in accordance with the “rule of law” that validate and verify the herein required tax status change to Affiant’s IRS files and records.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"There are only two types of people related to individual income tax liability: “taxpayers” and legal “non-taxpayers.” Affiant won't quote all of the many court case rulings on this distinction that you and the IRS should be well aware of and there is a well-established fact in law that Congress does not have “exclusive” legislative authority over the 50 states of the union and the living men and woman who live in these states.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"The U.S. Department of the Treasury defines the term “United States” in 31 U.S.C. 321 (d)(2) to mean only the National Government in the District of Columbia and not the 50 states of the Union per 26 U.S.C 7408(d). Title 26 excludes any reference to the Constitutional Republic where American National people live.", style=normal_style)
        docx_document.add_paragraph()

        notice_text = docx_document.add_paragraph("", style=centered_style)
        run = notice_text.add_run("NOTICE OF REVOCATION OF ELECTION")
        run.bold = True
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"26 U.S.C. 6013(g)(4) addresses Termination of Election and 6013(g)(4)(A) deals with Revocation of taxpayer, but this section is not the final governing law on this.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"{titlecased_name} as Affiant, is a living American National {gender} and not a “federal” “statutory” U.S. “person”, and to the best of {pronouns[1]} knowledge, does hereby expressly state {pronouns[1]} desire and intention to lawfully “terminate” and “revoke” {pronouns[1]} previous “presumed” (by the IRS) “election” to be treated as a federal citizen “taxpayer.” Effective immediately, Affiant has no further jurisdictional or contractual connection to the IRS or a federal “Individual” Income tax liability.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"Affiant is aware of and accepts the fact that {pronouns[0]} can never again make an “election” to become a “taxpayer” in the future. Let the record be established by {pronouns[1]} testimony herein for the express purpose of {pronouns[1]} Revocation of Election.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"As an American National with no “statutory” connections to the “foreign” jurisdiction of the IRS, and with the freedom from being compelled into “involuntary servitude” to the statutory “United States” [D.C.], no federal government may determine {pronouns[1]} “domicile” as that would amount to compelled association in direct violation of the Foreign Sovereign Immunities Act as per 28 U.S.C. 1605(a)(2) and the Thirteenth Amendment related to “slavery” and more importantly, {pronouns[1]} “inalienable rights” to exercise {pronouns[1]} free will to choose {pronouns[1]} legislative, political, and jurisdictional affiliations.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"Affiant does not derive any income that is effectively connected with the conduct of a statutory “trade or business” within the District of Columbia and {pronouns[0]} has no physical or statutory domicile within the jurisdiction of the District of Columbia.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"The limited and localized “municipal law” of the foreign owned and controlled municipal government in the District of Columbia, a foreign enclave state to {pronouns[1]} domicile, has no power or legal authority to force or compel a state Citizen or an American National to be domiciled within D.C.’s IRS jurisdiction unless they knowingly and willingly contract with it and consent to the IRS’s jurisdiction.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"Affiant, as an American National man, revokes, cancels, and makes void, ab initio, {pronouns[1]} power of appointment on any and all contracts and agreements, forms, or any instrument which may be construed in any way to give any federal IRS related agency any authority or jurisdiction over {pronouns[2]} unless {pronouns[0]} has first harmed or trespassed upon their rights. Brady v. U.S., 379 U.S. 742 at 748 (1970):", style=normal_style)
        docx_document.add_paragraph()

        quoted_text = docx_document.add_paragraph("", style=indent_list_level_1)
        run = quoted_text.add_run("“Waivers of Constitutional [protected] rights not only must be voluntary, they must be knowingly intelligent acts, done with sufficient “full disclosure” awareness of the relevant circumstances and consequences.”")
        run.italic = True
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"{titlecased_name} hereby revokes, rescinds, and makes void ab initio, all powers of attorney in fact or otherwise, implied in law or “presumed” by the IRS to exist that are related to the social security number the government assigned to a nom de guerre name spelled similar to {pronouns[1]} except in all capital letters (referring to an “artificial person”). The all capital letter spelled name, similar to {pronouns[1]} birth name, is not {pronouns[2]}, the living breathing man. As a living man, {pronouns[0]} does not have a social security number and {pronouns[1]} name is not spelled using all capital letters.", style=normal_style)
        docx_document.add_paragraph()

        paragraph = docx_document.add_paragraph("“In the United States of America, there are two (2) separated and distinct jurisdictions, such being the jurisdiction of the states within their own state boundaries, and the other being federal jurisdiction (United States), which is limited to the District of Columbia, the U.S. Territories, and federal enclaves within the states, under Article 1, Section 8, Clause 17,” ", style='Normal')
        run = paragraph.add_run("Bevans v. United States, 16 U.S. 336 (1818)")
        run.italic = True
        docx_document.add_paragraph()        

        docx_document.add_paragraph(f"State: The term “State” shall be construed to include the District of Columbia, where such construction is necessary to carry out provisions of this title.” 26 U.S.C sec. 7701. [with limited applicability to only federal citizens].", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"United States: The term “United States” when used in a geographical sense includes [is limited to] only the States [the District of Columbia and other federal territories within the borders of the states] and the District of Columbia, 26 U.S.C sec. 7701.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"According to RRA98 and the IRM, your Office is required to answer reasonable questions within 30 days. There is only one question that requires a written response from your Office;", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"Is there anything in the three (3) documents enclosed (including this one) that you don't understand or would like additional information or clarification on pertaining to Affiant’s present “foreign” jurisdiction to the IRS’s jurisdiction and {pronouns[1]} non-taxable status? Your silence will act as your and the IRS’s acceptance of Affiant’s non-taxable status.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"{titlecased_name} looks forward to your response to this question so there is no misunderstanding as to Affiant’s established non-taxable tax status.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"Please forward these Revocation of Election documents and instructions herein, if necessary or required, to appropriate IRS management, operational database personnel or to the IRS’s Chief Legal counsel.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"This Revocation of Election notice is effective immediately upon your receipt as agents within the Internal Revenue Service, the Internal Revenue, and the IRS.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"Congress in D.C. and all IRS related agencies, by statutory requirement must now recognize and acknowledge {titlecased_name}’s lawful “non taxpayer” (1040) status.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"The IRS records and databases must reflect this Revocation of Election as it has been legally established. All previous “presumed” to be taxable “elections” by the IRS have been duly “terminated” with your receipt of this Revocation of Election.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"Please make all necessary changes to the IRS’s databases to expedite Affiants’s non-taxable tax status.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"Sincerely,", style=normal_style)
        docx_document.add_paragraph(f"{titlecased_name}, a living {gender}", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"Attachments:", style=normal_style)

        #########################################################################################################
        #               START BULLET POINTS FOR ATTACHMENT SECTION
        # We have to do this in a funky way because there is not a native way to reset the counter for these bullet
        # points. This is a known issue when programmatically working with MS Word docs.
        #########################################################################################################
        # Add the first item and make it a numbered list item starting at #1 with less indentation
        paragraph1 = docx_document.add_paragraph(f"AFFIDAVIT", style='List Number 2')
        paragraph1._element.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = 1  # Keep the same list ID
        paragraph1._element.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = 1  # Decrease the level

        # Add the second item and make it a numbered list item starting at #2 with less indentation
        paragraph2 = docx_document.add_paragraph(f"Supplemental Information", style='List Number 2')
        paragraph2._element.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = 1  # Keep the same list ID
        paragraph2._element.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = 1  # Decrease the level
        #########################################################################################################
        #               END BULLET POINTS FOR ATTACHMENT SECTION
        #########################################################################################################

        #########################################################################################################
        #               END OF LETTER OF INTENT
        #########################################################################################################

    #########################################################################################################
    #                      END OF create_letter_of_intent FUNCTION
    #########################################################################################################


    #########################################################################################################
    #                       START OF create_affidavit FUNCTION
    #########################################################################################################
    def create_affidavit(self, docx_document, first_given_name, middle_given_name, family_name, man_or_woman, selected_states, street_address, city, zip_code, mailing_state, social_security_number, irs_commissioner, local_irs_service_center_street_address, local_irs_service_center_city, local_irs_service_center_state_list, local_irs_service_center_zip, republic_of_birth, sojourn_states_list, mailing_address_state_list, republic_of_birth_list, notary_state_list, notary_county):

        #########################################################################################################
        #                       START OF format_paragraph_keywords FUNCTION
        #########################################################################################################
        def format_paragraph_keywords(paragraph, formatted_words, styles):
            '''
            Format specific keywords in a paragraph with the given styles.

            Args:
                paragraph: The paragraph to format.
                formatted_words (list): A list of words or phrases to format.
                styles (list): A list of formatting styles ('bold', 'italic', 'underline').

            Returns:
                None
            '''
            paragraph_text = paragraph.text
            new_runs = []

            for word in formatted_words:
                index = paragraph_text.find(word)
                if index >= 0:
                    before = paragraph_text[:index]
                    word_to_format = paragraph_text[index:index + len(word)]
                    after = paragraph_text[index + len(word):]

                    new_runs.append((before, ''))
                    new_runs.append((word_to_format, styles))
                    paragraph_text = after

            new_runs.append((paragraph_text, ''))

            paragraph.clear()

            for run_text, run_styles in new_runs:
                run = paragraph.add_run(run_text)
                for style in run_styles:
                    if 'bold' in style:
                        run.bold = True
                    if 'italic' in style:
                        run.italic = True
                    if 'underline' in style:
                        run.underline = True
        #########################################################################################################
        #                       END OF format_paragraph_keywords FUNCTION
        #########################################################################################################

        #########################################################################################################
        #                       START OF DEFINING VARIABLES
        #########################################################################################################
        # Extract the "State Name" from each dictionary
        soujourn_state_names_list = [state_dict["State Name"] for state_dict in sojourn_states_list]
        soujourn_state_names = ", ".join(soujourn_state_names_list)

        # Extract the "State Abbreviations" from each dictionary
        soujourn_state_abbreviations_list = [state_dict["State Abbreviation"] for state_dict in sojourn_states_list]
        soujourn_state_abbreviations = ", ".join(soujourn_state_abbreviations_list)

        # Extract the "State " from each dictionary
        soujourn_republic_name_list = [state_dict["Republic Name"] for state_dict in sojourn_states_list]
        soujourn_republic_names = ", ".join(soujourn_republic_name_list)

        mailing_address_republic_name = mailing_address_state_list[0]['Republic Name']
        mailing_address_state_name = mailing_address_state_list[0]['State Name']
        mailing_address_state_abbreviation = mailing_address_state_list[0]['State Abbreviation']

        republic_of_birth_name = republic_of_birth_list[0]['Republic Name']
        republic_of_birth_state = republic_of_birth_list[0]['State Name']
        republic_of_birth_state_abbreviation = republic_of_birth_list[0]['State Abbreviation']

        local_irs_service_center_republic_name = local_irs_service_center_state_list[0]['Republic Name']
        local_irs_service_center_state__name = local_irs_service_center_state_list[0]['State Name']
        local_irs_service_center_state_abbreviation = local_irs_service_center_state_list[0]['State Abbreviation']

        notary_republic_name = notary_state_list[0]['Republic Name']
        notary_state_name = notary_state_list[0]['State Name']
        notary_state_abbrivation = notary_state_list[0]['State Abbreviation']

        vessel_name = first_given_name + ' ' + middle_given_name + ' ' + family_name
        live_name = first_given_name + '-' + middle_given_name + ': ' + family_name

        # Determine the pronouns based on the 'are_you' variable
        if man_or_woman == 'Man':
            pronouns = ("he", "his", "him")
            gender = man_or_woman.lower()
        elif man_or_woman == 'Woman':
            pronouns = ("she", "her", "her")
            gender = man_or_woman.lower()

        # Create variables for titlecased and uppercase versions of the name
        titlecased_name = vessel_name.title()
        uppercase_name = vessel_name.upper()

        # Get the current year
        current_year = datetime.now().year

        # Get the current date in the desired format
        current_date = datetime.now().strftime("%m/%d/%Y")
        #########################################################################################################
        #                       END OF DEFINING VARIABLES
        #########################################################################################################

        #########################################################################################################
        #               START OF CREATING HEADER
        #########################################################################################################
        # Create a header for all pages
        section = docx_document.sections[0]
        header = section.header

        # Add the header text
        header_text = f"Revocation of Election for {titlecased_name}"
        header_paragraph = header.paragraphs[0]
        header_run = header_paragraph.add_run(header_text)
        # Adjust the font size (e.g., set it to 16 points)
        header_run.font.size = Pt(12)
        header_run.bold = True
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_paragraph.space_after = Pt(12)  # Adjust the spacing as needed
        #########################################################################################################
        #               END OF CREATING HEADER
        #########################################################################################################

        #########################################################################################################
        #               START OF CREATING CUSTOM DOCUMENT STYLES
        #########################################################################################################

        # Create a custom style for normal text with a custom leading (line spacing)
        normal_style = docx_document.styles['Normal']
        normal_style.font.size = Pt(10)
        normal_style.paragraph_format.space_after = Pt(0)

        # Define a custom style with a left indent of 36 points
        indent_list_level_1 = docx_document.styles.add_style('IndentListLevel1', WD_STYLE_TYPE.PARAGRAPH)
        indent_list_level_1.font.size = Pt(10)
        indent_list_level_1.paragraph_format.space_after = Pt(0)
        indent_list_level_1.paragraph_format.left_indent = Pt(36)

        # Define a custom style with a left indent of 44 points
        indent_list_level_2 = docx_document.styles.add_style('IndentListLevel2', WD_STYLE_TYPE.PARAGRAPH)
        indent_list_level_2.font.size = Pt(10)
        indent_list_level_2.paragraph_format.space_after = Pt(0)
        indent_list_level_2.paragraph_format.left_indent = Pt(44)


        # Create a custom style for centered text
        centered_style = docx_document.styles.add_style('Centered', WD_STYLE_TYPE.PARAGRAPH)
        centered_style.font.size = Pt(10)
        centered_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        centered_style.paragraph_format.space_after = Pt(0)
        #########################################################################################################
        #              END OF CREATING CUSTOM DOCUMENT STYLES
        #########################################################################################################


        #########################################################################################################
        #               START OF LETTER OF AFFIDAVIT
        #########################################################################################################
        notice_text = docx_document.add_paragraph("", style=centered_style)
        run = notice_text.add_run("AFFIDAVIT")
        run.bold = True
        docx_document.add_paragraph("In support and validation of Affiant’s Revocation of Election rights", style=centered_style)
        docx_document.add_paragraph()


        docx_document.add_paragraph("KNOW ALL MEN BY THESE PRESENTS:", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"A change in Affiant’s Form 1040 federal Individual Income Tax status “election” notice to the IRS must be expressed in writing to be made of record, and in commerce, truth is sovereign and expressed in the form of an AFFIDAVIT, which in commercial law becomes the judgment of fact in commerce if not rebutted.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"I, {titlecased_name}, a living, breathing, sentient, flesh and blood {gender} and Affiant herein lawfully declare, without (not within) the statutory United States, the following:", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"Notice to principal is notice to agent, notice to agent is notice to principal and DEFAULT of principal is DEFAULT to agent.", style=normal_style)
        docx_document.add_paragraph()

        # REASON #1
        paragraph = docx_document.add_paragraph(f"This AFFIDAVIT is made as a matter of record of Affiant’s own right, sui juris, in {pronouns[1]} own propria persona status. Affiant, related to the IRC, does not consent to being classified as or referred to as a “legal fiction” trust or corporation entity or an artificial “construct” created by a foreign municipal government in the District of Columbia (D.C.), without (not within) Affiant’s jurisdiction and herein revokes and denies all “presumptions” by the IRS that Affiant is a legal fiction, or deceased, incompetent, in “probate,” lost at sea, or legally connected to the IRS’s jurisdiction in Washington, D.C. or Puerto Rico or anywhere else in the world.", style='List Number 2')
        format_paragraph_keywords(paragraph, ["sui juris", "propria persona"], ["italic"])
        docx_document.add_paragraph()
        
        # REASON #1
        docx_document.add_paragraph(rf"Affiant was born in the {republic_of_birth} and is over twenty-one, of sound mind, competent and able to testify in matters set forth herein, and Affiant has personal knowledge of the facts stated herein and waives none of {pronouns[1]} inalienable rights.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #3
        docx_document.add_paragraph(f"Affiant was born of white parents who were both born in Republic states of the union and Affiant, a living flesh and blood {gender}, is presently sojourning within the {soujourn_republic_names}, parts of the union of Republics, under the laws of God and “the rule of law” and then under the original Constitution (1787) and the Bill of Rights ratified in 1791, and does not waive any of {pronouns[1]} inalienable rights endowed by {pronouns[1]} Creator.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #4
        docx_document.add_paragraph(f"Affiant is not an “illegal tax protester” as {pronouns[0]} discharges all income taxes lawfully owed within {pronouns[1]} jurisdiction. Congress excluded the 50 union states from the IRS’s definition of “United States” for purposes of Subtitle A taxes and all “non federal” income from these 50 states is defined as being from “sources” without (not within) the IRS’s definition of the “United States” (IRC sec. 864(c)(4)(A).", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #5
        docx_document.add_paragraph(f"Affiant’s domicile and union state Citizenship as per 26 U.S.C. 865(g)(B) has no known lawful affiliation or connection to D.C.’s “municipal” government’s Fourteenth and Sixteenth Amendments and Affiant is “alien” or “foreign” to the “United States” as defined in the Internal Revenue Code (IRC) as the 10-mile square land area more commonly known as the District of Columbia.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #6
        docx_document.add_paragraph(f"Affiant is not a “U.S. person” or a “U.S. citizen” or a “taxpayer” as defined in 26 U.S.C. 7701 et seq. 26 U.S.C. is not “positive law” applicable to state Citizens of the union or to American Nationals as it is not promulgated in the Federal Register, thus, not applicable to Affiant as a living {gender} not “residing” or “domiciled” in D.C. or one of its territories or 'federal zones' within the 50 states of the union.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #7
        docx_document.add_paragraph(f"Affiant is not a “resident” or an “inhabitant” or a “franchise” or a “subject” or a “ward” or a “property” or a “chattel” of the jurisdiction of the “municipal” United States and the IRS, as {pronouns[0]} was not born or naturalized in the “United States” and thus, is not a “citizen” of the United States - defined as the District of Columbia.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #8
        docx_document.add_paragraph(f"Affiant is not a federal government “employee” and does not have a federal government “employer” and is not a “public officer” of the federal government in D.C. and is not a “U.S. resident alien” or a Fourteenth or Sixteenth Amendment affiliated “person” subject to the IRS’s prima facie presumptive “color of law” regulations, according to the Internal Revenue Code (hereinafter IRC).", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #9
        docx_document.add_paragraph(f"Affiant is not an “Officer” or “Agent” of a corporation under a “duty” to “withhold” from employees, or an elected Official or “Fiduciary” of the United States [D.C.] or any federal military base or private police force connected to government in Washington D.C., Puerto Rico, U.S. Virgin Islands, Guam or any other territory or possession of the United States [D.C].", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #10
        docx_document.add_paragraph(f"Affiant is not required to have a U.S. Tax I.D. number nor is {pronouns[0]} engaged in a “trade or business” as defined in the IRC. Affiant is not under the jurisdiction of the Federal Retirement Thrift Savings Plan.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #11
        docx_document.add_paragraph(f"As a living {gender} and not born in or on a federal territory and “non-resident” with respect to the “federal” United States located in D.C., Affiant’s “non-taxable” tax status qualifies {pronouns[2]} to demand that {pronouns[1]} Revocation of Election herein is duly noted by the IRS Commissioner and the DIRECTOR and all other IRS agents, and that Affiant’s IRS files and records within the IRS’s databases are changed to show Affiant’s “non-taxable” status, effective immediately upon your receipt.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #12
        docx_document.add_paragraph(f"Affiant reserves all of {pronouns[1]} Inalienable Rights granted by {pronouns[1]} Creator to be free from being forced against {pronouns[1]} will to comply with any foreign federal government or IRS adhesion contracts not fully disclosed before presumed by the IRS to be in force or presumed by the IRS that Affiant’s tax status is something different than “non-taxable” as per Affiant’s Revocation of Election notice herein.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #13
        docx_document.add_paragraph(f"Affiant does not derive any “gross income” or “taxable income” from sources within the United States [D.C.] and any private compensation Affiant might receive will be a non-taxable exchange for Affiant’s labor on an equal value exchange basis with no profit or gain from capital occurring.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #14
        docx_document.add_paragraph(f"Affiant will not accept federal-sourced “privileges” upon which an “excise” tax may be imposed.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #15
        docx_document.add_paragraph(f"Affiant sojourns within the {soujourn_republic_names} and is not domiciled in any geographic area within the {soujourn_republic_names} that is within or part of any “federal” district or “federal” “zone” as referred to in the Buck Act.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #16
        docx_document.add_paragraph(f"Affiant declares that all previous tax status “elections” were made by the IRS against Affiant’s free will choice, under threat, duress, and coercion, and without the IRS acting on Affiant’s previous demands to change {pronouns[1]} tax status.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #17
        docx_document.add_paragraph(f"Affiant declares that all past taxable “elections” presumed to exist by the IRS are null and void from the date of the IRS Commissioner’s and DIRECTOR’s receipt of Affiant’s Revocation of Election herein.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #18
        docx_document.add_paragraph(f"Affiant is not a “legal” or “statutory” federal “person” or “individual” subject to the IRS’s “foreign” and “alien” D.C. jurisdiction which is foreign to Affiant’s jurisdiction. Affiant has never agreed to be the “surety” or “guarantor” for alleged income tax debts for the government created “fiction” named {uppercase_name}.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #19
        docx_document.add_paragraph(f"Affiant declares that the federal government’s foreign and municipal Fourteenth and Sixteenth Amendments only have jurisdiction within the District of Columbia and its territories and not within the jurisdiction of Affiant as a living man who sojourns within the {soujourn_republic_names} and not in a “Buck Act” federal zone like {soujourn_state_abbreviations}.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #20
        docx_document.add_paragraph(f"Affiant declares that the term “United States” (meaning D.C.) is foreign with respect to the 50 Republic states of the Union and it cannot migrate to another sovereignty, such as, the {soujourn_republic_names}, as per 19 CJS 883.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #21
        docx_document.add_paragraph(f"Affiant declares that 26 U.S.C. has never been “positive” law as it and the IRC do not have certain required “implementing regulations” for enforcement promulgated in the Federal Register which is a lawful requirement to have said Statute and 26 U.S.C apply to American Nationals or state Citizens of the union (qualified non income tax volunteers and without the IRS’s jurisdiction).", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #22
        paragraph = docx_document.add_paragraph(f"Affiant believes Cynthia J. Mills, U.S. Treasury, Disclosure Officer said:", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"“The IRC is not positive law, it is special law applicable to specific persons in the United States [D.C.] who choose to make themselves subject to the requirement of the special laws in the IRC by entering into an employment agreement with the U.S. Government.”", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“The IRC is not positive law, it is special law applicable to specific persons in the United States [D.C.] who choose to make themselves subject to the requirement of the special laws in the IRC by entering into an employment agreement with the U.S. Government.”"], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f" Affiant does not choose to make himself “subject to” any “special IRS laws” applicable to specific (federal) persons and {pronouns[0]} is prohibited by law from doing so as a living man and not one of those IRS “special persons.”", style=indent_list_level_1)
        docx_document.add_paragraph()

        # REASON #23
        docx_document.add_paragraph(f"Affiant is not engaged in the manufacturing of Alcohol, Tobacco or Firearms or Energy conserving components or renewable energy sources.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #24
        docx_document.add_paragraph(f"Affiant does not consent to any “presumptions” by the IRS to be classified as a “taxpayer” within the federal government’s foreign constructive trust operations, including any federal policies that may fall under Martial law or Admiralty Law.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #25
        paragraph = docx_document.add_paragraph(f"Affiant has the free will inalienable right to declare {pronouns[1]} political and law forum jurisdiction and “tax status” and not to be treated like a Fourteenth and Sixteenth Amendment federal “person.” Said Amendments convert private law Merchant to be moved into the public sector that may control “public policy” and cause a loss to Affiant’s civil commercial common rights under Public Law Merchant as preserved in the court case of Swift v. Tyson.", style='List Number 2')
        format_paragraph_keywords(paragraph, ["Swift v. Tyson"], ["italic"])
        docx_document.add_paragraph()

        # REASON #26
        docx_document.add_paragraph(f"Affiant does not operate or “reside” in one of the “several states” of the territorial government within D.C.’s limited Article 1 legislative court jurisdiction and Affiant’s mailing location is non-domestic and zip code exempt.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #27
        docx_document.add_paragraph(f"Affiant herein “revokes” {pronouns[1]} consent to the IRS’s previous presumed taxable “election” of Affiant and to be under any form of Admiralty, Maritime, Military, private, or municipal law “in rem” concerning Affiant’s declared non-taxable American National status, as related to the IRC and any future imposed income taxes not applicable to {pronouns[1]} current non-taxable status and jurisdiction.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #28
        docx_document.add_paragraph(f"Affiant declares that there is no known bilateral contract or conveyance between Affiant and the IRS “in personam” or in “venue” or in “subject matter” establishing the “res” or the “thing” of the contract subject to IRS’s Roman Civil Law constructs and without (not within) Affiant’s jurisdiction.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #29
        docx_document.add_paragraph(f"Affiant had not been born in 1933 when the U.S. went off the gold standard (law) thus, Affiant was not a part of that major commercial agreement change in law forum. Thus, the IRS has no “in rem” jurisdiction to compel performance on a limited “unilateral” contract against Affiant when Affiant was not a party to that contract in 1933 (House Joint Resolution 192), Public Law 73-10, and does not wish to exercise {pronouns[1]} option to be part of or under the government’s “public policy” statutory law and constructive trust operations or Martial Law related thereto.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #30
        docx_document.add_paragraph(f"Affiant, acting in “good faith” and in honor, declares that {pronouns[0]} qualifies as a living man and has the right to “revoke” the previous “taxpayer” election presumed by the IRS. Affiant authorizes and demands that the IRS Commissioner and the DIRECTOR or their delegates to make said changes to all of Affiant’s income tax records and IRS files effective immediately. 26 U.S.C 6013 (g)(4)(A) addresses the right to “revoke” a previous tax status “election,” however, Affiant does not require or need 26 U.S.C. 6013 when exercising {pronouns[1]} superior law forum, being {pronouns[1]} “inalienable right” to choose {pronouns[1]} taxing authority and {pronouns[1]} taxing jurisdiction.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #31
        docx_document.add_paragraph(f"Affiant is not now nor ever intends to be a “registered” voter in “federal” elections and thus, is not receiving the “privilege” of voting in federal elections. As a living man and not a federal “U.S.citizen”, Affiant has no legal right to vote in any federal elections that are in a 'foreign' D.C. jurisdiction to {pronouns[1]} jurisdiction.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #32
        docx_document.add_paragraph(f"Affiant has never been a federal (public official) jurist in a trial and as a non-federal living man, Affiant cannot act in a “federal court” in a public capacity.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #33
        docx_document.add_paragraph(f"Affiant is not a member of a 501(c)(3) Church and does not receive the “benefit” of a tax deduction for donations made and Affiant does not take “tax deductions.”", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #34
        docx_document.add_paragraph(f"Every variant of a compelled government-issued “benefit,” “license” or “privilege” which Affiant has been forced to use to exist in today’s society that may attempt to identify Affiant as either a “trustee,” “fiduciary,” “public officer,” “public official” or a “federal citizen” or “US citizen,” is against Affiant’s free will choice to be without (not within) the “federal” meaning of said “labels” and IRS “terms” and are herein expressly rebutted and denied by Affiant.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #35
        paragraph = docx_document.add_paragraph(f"Affiant, under {pronouns[1]} God given inalienable rights and {pronouns[1]} American National foreign jurisdiction to the IRS's, does not consent to having {pronouns[1]} name unlawfully “converted” to an all capital letter nom-de-guerre spelling (which is not Affiant’s organic name received at birth) to create an unwanted foreign law “Res” connection to D.C. and the IRS’s foreign jurisdiction to Affiant’s jurisdiction.", style='List Number 2')
        format_paragraph_keywords(paragraph, ["nom-de-guerre"], ["italic"])
        docx_document.add_paragraph()

        # REASON #36
        docx_document.add_paragraph(f"Affiant herein revokes all further consent to be a “Res - ident” (the “thing” identified) juristic statutory servant “individual” resident to a Master federal government under Roman Civil municipal law that is foreign to Affiant’s American National status and {pronouns[1]} superior law forum as a living man.", style='List Number 2')
        docx_document.add_paragraph()
        
        # REASON #37
        docx_document.add_paragraph(f"Affiant revokes herein ab initio and rebuts any signature and connection to Form SS-5 and voter registration, to the extent said relationships connect Affiant to the IRS’s federal D.C. jurisdiction. Affiant does not consent to be a Federal Trustee in the Federal government’s constructive trust “public policy” agenda and Affiant has no known foreign earned income earned within the United States [D.C.].", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #38
        docx_document.add_paragraph(f"Affiant does not consent to be classified as a federal “Debtor” in any pending or future United States [D.C.] bankruptcy or Martial Law proceedings.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #39
        docx_document.add_paragraph(f"Affiant is not an “enemy” or an “enemy combatant” against the federal government or a political or financial “terrorist” planning or scheming against the United States or the IRS and is not involved or active in any such group. Affiant is only interested in asserting {pronouns[1]} right to be classified as “non-taxable” under the laws passed by Congress under the original Constitution that protect American Nationals and state Citizens from taxes being unlawfully imposed upon Affiant.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #40
        docx_document.add_paragraph(f"Affiant’s (Form 1040) income tax status change demand to “non-taxable” with {pronouns[1]} Revocation of Election herein, is a change in the IRS’s previously presumed tax status of Affiant and thus, Affiant has the right to receive a confirmation of your receipt and execution of Affiant’s changed tax status to non-taxable.", style='List Number 2')
        docx_document.add_paragraph()


        docx_document.add_paragraph(f"Affiant’s non-taxable “foreign jurisdiction” to the IRS’s D.C. or Puerto Rico jurisdiction has been well established here with this ROE and AFFIDAVIT and does not require the approval of the IRS. Affiant’s notice of {pronouns[1]} non-taxable status with {pronouns[1]} ROE documents herein, is the truth in law unless rebutted in writing by a high ranking senior IRS officer.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"Should you or any other officer within the Office of the Commissioner of the IRS or DIRECTOR’s Office wish to rebut this AFFIDAVIT, said rebuttal must be from either you or one of your high ranking authorized delegates in their individual and personal capacity under a sworn Affidavit of Rebuttal, executed and personally signed and notarized as being true, correct and complete.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"Any IRS disagreement rebuttal of Affiant’s declarations herein and {pronouns[1]} right to be an American National not “subject to” the IRC and the IRS’s Subtitle A individual income taxes, must be expressed to Affiant in writing within the time frame allowed by law.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"As the IRS Commissioner and DIRECTOR with a “duty” to perform as instructed herein and in your individual capacity as a living man (or woman), Affiant believes you have 30 days under RRA98 to reply to Affiant with an acknowledgment that {pronouns[1]} tax status has been changed to reflect {pronouns[1]} non-taxable status.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"In the event you or your delegate does not acknowledge Affiant’s tax status change in a timely manner as required per RRA98 and other legal precedents, you and your Office will not be in compliance with RRA98. Your silence will be construed by Affiant to be your tacit approval of Affiant’s Revocation of Election, {pronouns[1]} AFFIDAVIT, and {pronouns[1]} non-taxable status.", style=normal_style)
        docx_document.add_paragraph()

        paragraph = docx_document.add_paragraph(f"“Silence gives consent, is the rule of business life. Express consent, then, not being necessary, is there anything from which consent may be implied? There is length of time.” Padelford, Fay & Co. v. Mayor and Alderman of City of Savannah, 14 Ga. 438,WL. 1492, (1854).", style=indent_list_level_1)
        format_paragraph_keywords(paragraph, ["“Silence gives consent, is the rule of business life. Express consent, then, not being necessary, is there anything from which consent may be implied? There is length of time.” Padelford, Fay & Co. v. Mayor and Alderman of City of Savannah"], ["italic"])
        docx_document.add_paragraph()

        paragraph = docx_document.add_paragraph(f"“Silence” is species of conduct, and constitutes an implied representation of existence of facts in question... where silence is of such character and under such circumstances that it would become fraud... it will operate as estoppel. Carmine v. Bowen 64 AT 932.", style=normal_style)
        format_paragraph_keywords(paragraph, ["Carmine v. Bowen"], ["italic"])
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"Affiant hereby affixes {pronouns[1]} autograph to all above declarations and statements with explicit reservations of all of {pronouns[1]} inalienable rights and without prejudice to any of those rights;", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"Affiant does hereby state the following: I, {titlecased_name}, declare under penalty of perjury under the laws of the United States of America, without the “United States” [District of Columbia] that the foregoing is true and correct in fact and in substance, to the best of My current information, knowledge, and belief, per 28 U.S.C.1746 (1).", style=normal_style)
        docx_document.add_paragraph()

        # Insert a page break
        run = docx_document.add_paragraph().add_run()
        run.add_break(WD_BREAK.PAGE)

        docx_document.add_paragraph(f"FURTHER, THIS AFFIANT SAYETH NOT ON THIS __________ DAY OF _______________ {current_year}", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"Executed at __________________, {notary_republic_name}", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"by__________________________  Date____________________ ", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"{titlecased_name}", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"A living {gender}, in Honor and in good faith by “special appearance” with all unalienable rights reserved.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"{titlecased_name}, Authorized Representative", style=normal_style)
        docx_document.add_paragraph(f"c/o Non-domestic", style=normal_style)
        docx_document.add_paragraph(f"{street_address}", style=normal_style)
        docx_document.add_paragraph(f"{city}, {mailing_state}", style=normal_style)
        docx_document.add_paragraph(f"USA [near {zip_code}]", style=normal_style)

        docx_document.add_paragraph()

        docx_document.add_paragraph(f"SS #: {social_security_number}", style=normal_style)

        paragraph = docx_document.add_paragraph(f"assigned by a “federal” government agency to {uppercase_name}, nom-de-guerre, vessel, securitized negotiable instrument and not assigned to {titlecased_name}, a living man.", style=normal_style)
        format_paragraph_keywords(paragraph, ["nom-de-guerre"], ["italic"])

        docx_document.add_paragraph()

        docx_document.add_paragraph(f"{notary_republic_name}", style=normal_style)
        docx_document.add_paragraph(f"County of {notary_county}", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"On this day Affiant came before me and declared the signature / autograph is true and complete on the foregoing Affidavit. {titlecased_name} the undersigned, upon proper identification, personally came before me, a notary public and duly declared the truth of the foregoing with Affidavit in my presence. The Affidavit also acknowledged the autographing thereof to be {pronouns[1]} own voluntary act and deed.", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"Autographed this ___________ day of ______________________, {current_year} at ______________________", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"My Commission expires on___________________", style=normal_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph(f"By_____________________________________         Seal", style=normal_style)
        docx_document.add_paragraph()
        #########################################################################################################
        #               END OF LETTER OF AFFIDAVIT
        #########################################################################################################

    #########################################################################################################
    #                       END OF create_affidavit FUNCTION
    #########################################################################################################


    #########################################################################################################
    #                       START OF create_supporting_evidence FUNCTION
    #########################################################################################################
    def create_supporting_evidence(self, docx_document, first_given_name, middle_given_name, family_name, man_or_woman, selected_states, street_address, city, zip_code, mailing_state, social_security_number, irs_commissioner, local_irs_service_center_street_address, local_irs_service_center_city, local_irs_service_center_state_list, local_irs_service_center_zip, sojourn_states_list, mailing_address_state_list, republic_of_birth_list, notary_state_list, notary_county):
        
        #########################################################################################################
        #                       START OF format_paragraph_keywords FUNCTION
        #########################################################################################################        
        def format_paragraph_keywords(paragraph, formatted_words, styles):
            '''
            Format specific keywords in a paragraph with the given styles.

            Args:
                paragraph: The paragraph to format.
                formatted_words (list): A list of words or phrases to format.
                styles (list): A list of formatting styles ('bold', 'italic', 'underline').

            Returns:
                None
            '''
            paragraph_text = paragraph.text
            new_runs = []

            for word in formatted_words:
                index = paragraph_text.find(word)
                if index >= 0:
                    before = paragraph_text[:index]
                    word_to_format = paragraph_text[index:index + len(word)]
                    after = paragraph_text[index + len(word):]

                    new_runs.append((before, ''))
                    new_runs.append((word_to_format, styles))
                    paragraph_text = after

            new_runs.append((paragraph_text, ''))

            paragraph.clear()

            for run_text, run_styles in new_runs:
                run = paragraph.add_run(run_text)
                for style in run_styles:
                    if 'bold' in style:
                        run.bold = True
                    if 'italic' in style:
                        run.italic = True
                    if 'underline' in style:
                        run.underline = True
        #########################################################################################################
        #                       END OF format_paragraph_keywords FUNCTION
        #########################################################################################################

        #########################################################################################################
        #                       START OF DEFINING VARIABLES
        #########################################################################################################
        # Extract the "State Name" from each dictionary
        soujourn_state_names_list = [state_dict["State Name"] for state_dict in sojourn_states_list]
        soujourn_state_names = ", ".join(soujourn_state_names_list)

        # Extract the "State Abbreviations" from each dictionary
        soujourn_state_abbreviations_list = [state_dict["State Abbreviation"] for state_dict in sojourn_states_list]
        soujourn_state_abbreviations = ", ".join(soujourn_state_abbreviations_list)

        # Extract the "State " from each dictionary
        soujourn_republic_name_list = [state_dict["Republic Name"] for state_dict in sojourn_states_list]
        soujourn_republic_names = ", ".join(soujourn_republic_name_list)

        mailing_address_republic_name = mailing_address_state_list[0]['Republic Name']
        mailing_address_state_name = mailing_address_state_list[0]['State Name']
        mailing_address_state_abbreviation = mailing_address_state_list[0]['State Abbreviation']

        republic_of_birth_name = republic_of_birth_list[0]['Republic Name']
        republic_of_birth_state = republic_of_birth_list[0]['State Name']
        republic_of_birth_state_abbreviation = republic_of_birth_list[0]['State Abbreviation']

        local_irs_service_center_republic_name = local_irs_service_center_state_list[0]['Republic Name']
        local_irs_service_center_state__name = local_irs_service_center_state_list[0]['State Name']
        local_irs_service_center_state_abbreviation = local_irs_service_center_state_list[0]['State Abbreviation']

        notary_republic_name = notary_state_list[0]['Republic Name']
        notary_state_name = notary_state_list[0]['State Name']
        notary_state_abbrivation = notary_state_list[0]['State Abbreviation']

        vessel_name = first_given_name + ' ' + middle_given_name + ' ' + family_name
        live_name = first_given_name + '-' + middle_given_name + ': ' + family_name

        # Create variables for titlecased and uppercase versions of the name
        titlecased_name = vessel_name.title()
        uppercase_name = vessel_name.upper()

        # Determine the pronouns based on the 'are_you' variable
        if man_or_woman == 'Man':
            pronouns = ("he", "his", "him")
            gender = man_or_woman.lower()
        elif man_or_woman == 'Woman':
            pronouns = ("she", "her", "her")
            gender = man_or_woman.lower()

        # Get the current date in the desired format
        current_date = datetime.now().strftime("%m/%d/%Y")
        #########################################################################################################
        #                       END OF DEFINING VARIABLES
        #########################################################################################################

        #########################################################################################################
        #               START OF CREATING HEADER
        #########################################################################################################
        # Create a header for all pages
        section = docx_document.sections[0]
        header = section.header

        # Add the header text
        header_text = f"Revocation of Election for {titlecased_name}"
        header_paragraph = header.paragraphs[0]
        header_run = header_paragraph.add_run(header_text)
        # Adjust the font size (e.g., set it to 16 points)
        header_run.font.size = Pt(12)
        header_run.bold = True
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_paragraph.space_after = Pt(12)  # Adjust the spacing as needed
        #########################################################################################################
        #               END OF CREATING HEADER
        #########################################################################################################

        #########################################################################################################
        #               START OF CREATING CUSTOM DOCUMENT STYLES
        #########################################################################################################

        # Create a custom style for normal text with a custom leading (line spacing)
        normal_style = docx_document.styles['Normal']
        normal_style.font.size = Pt(10)
        normal_style.paragraph_format.space_after = Pt(0)


        # Define a custom style with a left indent of 36 points
        indent_list_level_1 = docx_document.styles.add_style('IndentListLevel1', WD_STYLE_TYPE.PARAGRAPH)
        indent_list_level_1.font.size = Pt(10)
        indent_list_level_1.paragraph_format.space_after = Pt(0)
        indent_list_level_1.paragraph_format.left_indent = Pt(36)

        # Define a custom style with a left indent of 44 points
        indent_list_level_2 = docx_document.styles.add_style('IndentListLevel2', WD_STYLE_TYPE.PARAGRAPH)
        indent_list_level_2.font.size = Pt(10)
        indent_list_level_2.paragraph_format.space_after = Pt(0)
        indent_list_level_2.paragraph_format.left_indent = Pt(44)

        # Create a custom style for centered text
        centered_style = docx_document.styles.add_style('Centered', WD_STYLE_TYPE.PARAGRAPH)
        centered_style.font.size = Pt(10)
        centered_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        centered_style.paragraph_format.space_after = Pt(0)
        #########################################################################################################
        #               END OF CREATING CUSTOM DOCUMENT STYLES
        #########################################################################################################

        #########################################################################################################
        #               START OF LETTER OF SUPPLEMENTAL INFORMATION
        #########################################################################################################
        docx_document.add_paragraph("Supplemental Information validating Affiant’s Revocation of Election", style=centered_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph("100 REASONS WHY CERTAIN AMERICAN NATIONALS AND STATE", style=centered_style)
        docx_document.add_paragraph("CITIZENS ARE NOT LIABLE FOR PAYING A FEDERAL INCOME TAX", style=centered_style)
        docx_document.add_paragraph()

        docx_document.add_paragraph("The legal research on the subject of “who is” and “who is not” liable to file a Form 1040 related to Affiant includes the following:", style=normal_style)
        docx_document.add_paragraph()

        # REASON #1
        docx_document.add_paragraph(f"The definition of a “taxpayer” within the Internal Revenue Code (IRC) has its foundation deeply rooted in the matter of “jurisdiction,” and it can be simply stated that jurisdiction and one’s “Domicile” related to one’s tax liability is most important in law when it comes to “who is” and “who is not” a “taxpayer.”", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #2
        docx_document.add_paragraph(f"The 50 states of the union under the original Constitution for the United States of America (1787) are essentially “foreign jurisdictions” or “foreign states” with respect to each other and with respect to the “Federal” government’s seat in the 10-mile square land area located in and on the District of Columbia (D.C.).", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #3
        paragraph = docx_document.add_paragraph(f"It is equally well-settled in law that the 50 states of the union are to be considered, with respect to the Internal Revenue Code (herein after IRC) and 26 USC, to be foreign to each other and that the courts of one State are not presumed to know and therefore are not bound to take judicial notice of the laws of another state.  Hanley v. Donahue, 116 U.S. 1, 29 L.Ed 535 6 S. ct 242,244 (1885).", style='List Number 2')
        format_paragraph_keywords(paragraph, ["Hanley v. Donahue"], ["italic"])
        docx_document.add_paragraph()

        # REASON #4
        docx_document.add_paragraph(f"Another key U.S. Supreme Court authority on this “foreign” status matter is the case of In re: Merriam’s Estate, 36 N.E. 505 (1894). The author of Corpus Juris Secondum (CJS), a legal encyclopedia, relied in part upon this case to arrive at the following conclusion about the “foreign” corporate status of the Federal government:", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"“The United States government is a foreign corporation with respect to a state.” [citing In re: Merriam’s Estate (supra) affirmed U.S. v. Perkins 16 S.Ct. 1073, 163 U.S. 625, 41 L.Ed. 387]", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“The United States government is a foreign corporation with respect to a state.”", "In re: Merriam’s Estate (supra) affirmed U.S. v. Perkins"], ["italic"])
        docx_document.add_paragraph()

        # REASON #5
        paragraph = docx_document.add_paragraph(f"Black’s Law Dictionary, 6th Ed., clearly defines “foreign state” as follows:", style='List Number 2')
        format_paragraph_keywords(paragraph, ["Black’s Law Dictionary"], ["italic"])
        docx_document.add_paragraph()        
        paragraph = docx_document.add_paragraph(f"“The several United States are considered “foreign” to each other except as regards to their relations as common members of the Union … one State of the Union is foreign to another in the sense of that rule.”", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“The several United States are considered “foreign” to each other except as regards to their relations as common members of the Union … one State of the Union is foreign to another in the sense of that rule.”"], ["italic"])
        docx_document.add_paragraph()

        # REASON #6
        quoted_state_abbreviations = ", ".join([f'“{abbrev}”' for abbrev in soujourn_state_abbreviations_list])
        paragraph = docx_document.add_paragraph(f"The IRS and its affiliated offices in D.C. are in a venue that is “foreign” to Affiant’s as a living man living in the {soujourn_republic_names} and not domiciled in the {soujourn_state_names} “federal zone” according to the Buck Act and referred to as {quoted_state_abbreviations}. Thus, Affiant proceeds at all times with explicit reservation of all of {pronouns[1]} inalienable rights to due process of law, “without” (outside) D.C.’s “federal” statutory and “municipal” law jurisdiction.", style='List Number 2')
        docx_document.add_paragraph()        
        paragraph = docx_document.add_paragraph(f"“The laws of Congress in respect to those matters [outside of Constitutionally delegated powers] do not extend into the territorial limits of the states [of the Union], but have force only in the District of Columbia and other places that are within the exclusive jurisdiction of the national government.” Caha v. U.S., 152 U.S. 211 (1894).", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“The laws of Congress in respect to those matters [outside of Constitutionally delegated powers] do not extend into the territorial limits of the states [of the Union], but have force only in the District of Columbia and other places that are within the exclusive jurisdiction of the national government.”", "Caha v. U.S."], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"The U.S. Supreme Court case, Afroyim v. Rusk, 387 U.S. 253 (1967), is another decision which restricts the federal government from creating indentured servants when it stated:", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"“In the [constitutionally defined] United States the people are sovereign, and the government cannot sever its relationship to the people by taking away their citizenship.”", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“In the [constitutionally defined] United States the people are sovereign, and the government cannot sever its relationship to the people by taking away their citizenship.”"], ["italic"])
        docx_document.add_paragraph()

        # REASON #7
        paragraph = docx_document.add_paragraph(f"The original U.S. Constitution states at Article 1, Section 9, Clause 4, to wit:", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"“No Capitation, or other direct Tax shall be laid, unless in Proportion [apportioned] to the Census ....”", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“No Capitation, or other direct Tax shall be laid, unless in Proportion [apportioned] to the Census ....”"], ["italic"])
        docx_document.add_paragraph()

        # REASON #8
        paragraph = docx_document.add_paragraph(f"The original U.S. Constitution protects Affiant’s birthright to “life, liberty and the pursuit of happiness” [and property] which the founding father framers and U.S. Supreme Court have declared includes Affiant’s inalienable right to contract,  acquire, to sell, rent and exchange properties of various kinds without requesting or taking any “privilege” or “franchise” from government [D.C.].", style='List Number 2')
        format_paragraph_keywords(paragraph, ["“life, liberty and the pursuit of happiness”"], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Said inalienable rights include exchanging Affiant's labor property for other properties including financial instruments like Federal Reserve Notes (FRNs) notwithstanding the private and foreign nature of FRNs.", style=indent_list_level_1)
        docx_document.add_paragraph()

        # REASON #9
        paragraph = docx_document.add_paragraph(f"In Murdock v. Pennsylvania (1943), the U.S. Supreme Court stated:", style='List Number 2')
        format_paragraph_keywords(paragraph, ["Murdock v. Pennsylvania"], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"“A state [D.C. foreign state] may not impose a charge [tax] for the enjoyment of a right granted by the Constitution … that unalienable rights are rights against which no lien can be established.”", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“A state [D.C. foreign state] may not impose a charge [tax] for the enjoyment of a right granted by the Constitution … that unalienable rights are rights against which no lien can be established.”"], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f" [This quote however is slightly in error, as the Constitution did not grant anyone any rights. Un-alien-able rights come from the Creator who granted man and woman “Agency” and “Dominion” a.k.a. “sovereignty,” and the original Constitution’s main purpose is to “protect” the unalienable “Rights” granted to us by Creator.]", style=indent_list_level_1)
        docx_document.add_paragraph()

        # REASON #10
        paragraph = docx_document.add_paragraph(f"Affiant has never knowingly, willingly or voluntarily relinquished {pronouns[1]} inalienable rights 'status' granted to {pronouns[2]} by {pronouns[1]} Creator, to become adhesioned to the D.C.’s “federal” municipal jurisdiction, well-known to be “foreign” to {pronouns[2]}; and {pronouns[0]} was never provided sufficient “full disclosure” when {pronouns[0]} inadvertently and mistakenly was mislead by IRS agents as to {pronouns[1]} correct tax status.", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Affiant intends herein to rectify this incorrect tax status problem by lawfully removing himself from the IRS’s “municipal” jurisdiction when you make all appropriate and lawful changes in the IRS’s databases as per {pronouns[1]} Revocation of Election herein, effectively “revoking” {pronouns[1]} taxable status as previously “presumed” to exist by the IRS.", style=indent_list_level_1)
        docx_document.add_paragraph()

        # REASON #11
        paragraph = docx_document.add_paragraph(f"The Forms 1040 Affiant has mistakenly sent to the IRS contained no clear and  concise instructions or references or adequate disclosure (in bad faith by the IRS),  explaining who is and is not subject to Subtitle A income taxes. The IRS has done a masterful job of obfuscating and hiding the fact that American Nationals (most Americans in the 50 states) and “nonresident alien individuals” (meaning most state Citizens of the union with no “federal” jurisdiction connections) are not liable for federal “individual” income taxes imposed by Subtitle A of the IRC.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #12
        paragraph = docx_document.add_paragraph(f"In U.S. Supreme Court cases like Flint v. Stone Tracy Co. 220 U.S 107 (1911)  and Pollock v. Farmer’s Loan and Trust Co., 157 U.S. 429, (1895), the Federal government learned that the power they thought they had to tax state Citizens of the Union (a.k.a. American Nationals) was not authorized by standing decisions of  the U.S. Supreme Court or by the original Constitution with only 13 Amendments.", style='List Number 2')
        format_paragraph_keywords(paragraph, ["Flint v. Stone Tracy Co.", "Pollock v. Farmer’s Loan and Trust Co., 157 U.S. 429,"], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"The famous U.S. Supreme Court decision in Pollock v. Farmer’s Loan and Trust Co. is what influenced the passing of the so-called Sixteenth Amendment, and most importantly, the Pollock decision is one of the main reasons why the Sixteenth Amendment has limited applicability only to the “National Government” in the District of Columbia. As President William Taft stated in the  Congressional Record (see below), that Amendment mainly applies to people working for or who are connected to the federal government and who are domiciled within the 10-mile square area known as Washington, D.C.", style=indent_list_level_1)
        format_paragraph_keywords(paragraph, ["Pollock v. Farmer’s Loan and Trust Co."], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"After the landmark decision in Pollock v. Farmer’s Loan and Trust Co. informed the “D.C.” federal government that taxing the income of state Citizens was unconstitutional, President William H. Taft made it quite clear that the legislative intent of the Sixteenth Amendment limited the IRS’s taxing authority and jurisdiction to the “National Government” in D.C. and its federal territories, possessions, and enclaves, and unquestionably it did not apply to state Citizens of  the union or American Nationals operating under the original Constitution.", style=indent_list_level_1)
        format_paragraph_keywords(paragraph, ["Pollock v. Farmer’s Loan and Trust Co."], ["italic"])
        docx_document.add_paragraph()

        # REASON #13
        paragraph = docx_document.add_paragraph(f"Furthermore, the Code of Federal Regulations at 26 CFR 1.871-1(a) and 26 U.S.C. 7701(b)(1)(B) use the term “nonresident alien individuals” to describe someone who IS NOT a taxable “U.S. person” or “U.S. citizen” subject to federal  municipal law. Certain federal municipal laws are only applicable to D.C.’s limited 10-mile square area jurisdiction and to its federal territories, possessions, and all other federal enclaves, i.e., geographic areas that are not states of the union. ", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #14
        paragraph = docx_document.add_paragraph(f"The term Affiant uses to describe {pronouns[1]} non-taxable tax status, “American National”, is similar to the IRC term “nonresident alien individual” (with no federal government connection resulting in a tax obligation), and thus, that term confirms that the IRC and the IRS acknowledge and admit to the fact that certain “natural-born” nonresident aliens (to D.C.), state Citizens and American Nationals  are “non-taxable” or “non-taxpayers.”", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Affiant's definition of the term “American National” related to {pronouns[1]} Revocation of Election herein means, a “non-statutory” living man who was born in one of the 50 states of the union, born from white parents and at least one of whom was born in  one of the 50 states, or who has been naturalized into the Constitutional Republic.  In this context, please refer to the Guarantee Clause in the U.S. Constitution.", style=indent_list_level_1)
        docx_document.add_paragraph()

        # REASON #15
        paragraph = docx_document.add_paragraph(f"By Affiant's birth and parentage and according to American Jurisprudence 2d., Sec. 2689 and 8 U.S.C. 1401(a), Affiant is not someone “within” (inside) the income taxing jurisdiction of the United States - defined as the District of Columbia per 26 U.S.C. 7408(d).", style='List Number 2')
        format_paragraph_keywords(paragraph, ["American Jurisprudence"], ["italic"])
        docx_document.add_paragraph()

        # REASON #16
        paragraph = docx_document.add_paragraph(f"Because Affiant's free will choice to choose {pronouns[1]} political affiliation rests solely  with Affiant, the burden of proof that {pronouns[0]} is a “federal,” “statutory,” “fiction,” “juristic,” “taxable,” “U.S. person,” “U.S. citizen,” “artificial person,” “inhabitant,” or a taxable “vessel” employed by or connected to the Federal Government and domiciled in D.C., falls 100% upon the IRS according to 5 U.S.C. However, 5 U.S.C.'s “Administrative” regulations do not apply to American Nationals. According to higher Court decisions and rulings which held that Affiant  has no lawful obligation to “prove a negative” - that {pronouns[0]} doesn't owe an income tax,  the “burden of proof” falls completely upon the IRS to prove that Affiant is a “taxpayer” and not an American National “non-taxpayer.” In this context, it is believed by Affiant that it is also correct that RRA98 Section 3001 has shifted the  burden of proof on to the Secretary of the Treasury or {pronouns[1]} delegate(s).", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #17
        paragraph = docx_document.add_paragraph(f"According to Federal Government records dating back to 1925 concerning the  Statutes at Large, including Title 26 of the U.S. Code, there is clear evidence that many Statutes at Large are “administrative” codes only, and they only apply to Fourteenth Amendment “federal” citizens. Cf. “Federal citizenship” in Black’s Law Dictionary, 6th Ed. A typical example is the Privacy Act at 5 U.S.C. 552a(a)(2), which expressly defines “individual” to include only “federal” citizens and resident aliens. Affiant is neither a “federal” citizen nor a “resident alien.”", style='List Number 2')
        format_paragraph_keywords(paragraph, ["Black’s Law Dictionary"], ["italic", "underline"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"This means that those Statutes at Large are basically “administrative restrictions” upon “federal citizens” and “government employees” whether or not said federal “persons” are aware of their contractually adhesioned relationship with the municipal jurisdiction of the federal government domiciled in Washington, D.C. These “administrative” Statutes have very little to do with Fundamental Laws established to protect the “inalienable” Rights of living men and woman who are not “federal citizens” or federal “resident aliens”.", style=indent_list_level_1)
        docx_document.add_paragraph()

        # REASON #18
        paragraph = docx_document.add_paragraph(f"Although Congress never incorporated the Federal Government, it did incorporate D.C. as a municipal corporation in 1871. That municipal corporation’s  stock, 100% of it, has been owned by entities of the “foreign” offshore International Monetary Fund (IMF) at least since the Bretton Woods Agreement was codified in 1944 at 22 U.S.C. 286 et. seq.", style='List Number 2')
        format_paragraph_keywords(paragraph, ["et. seq."], ["italic"])
        docx_document.add_paragraph()

        # REASON #19
        paragraph = docx_document.add_paragraph(f"Public Law 97-280 (96 Stat. 1211) declared the Holy Bible as the word of God. Affiant is quite certain that God wanted man to be ruled by His laws and the laws  of nature, and not to have man, with {pronouns[1]} lust for power, attempt to rule over His laws.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #20
        paragraph = docx_document.add_paragraph(f"Due to Affiant's free will inalienable Right to choose a political affiliation, {pronouns[0]} declares herein that {pronouns[0]} is not now, nor has {pronouns[0]} ever knowingly chosen or elected, to  be a Fourteenth Amendment “indentured” federal citizen, as {pronouns[0]} has no intention to  ever relinquish {pronouns[1]} “inalienable” God-given Rights in exchange for limited federal  government “privileges” that can be taken away from {pronouns[2]} at the federal government’s whim.", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Affiant's notice to you herein of {pronouns[1]} “non-statutory” American National status is {pronouns[1]} notice of {pronouns[1]} correct political and jurisdictional affiliation as it relates to {pronouns[1]} domicile “without” (outside) the IRS's “domestic” municipal jurisdiction.", style=indent_list_level_1)
        docx_document.add_paragraph()

        # REASON #21
        paragraph = docx_document.add_paragraph(f"The IRS’s statutory authority is “administrative” only and only applies to certain “types” of people (which Affiant is not) who “voluntarily” choose to “donate” their property (e.g., Federal Reserve Notes) to the “foreign” U.S. Treasury in D.C. or Puerto Rico or who are “federal persons” by consent or by “private municipal law” inside a Federal “domestic” jurisdiction.", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Title 31 of the U.S. Code says income taxes (not really taxes) paid to the U.S. Treasury are considered to be nothing more than “voluntary donations” to the Treasury of the United States [D.C.] by all those who “donate” their contributions with their 1040 Forms. There is no “mandatory” law compelling certain state Citizens or American Nationals to 'donate' their Federal Reserve Notes or other assets obtained from the “private sector” to the foreign-controlled Treasury that records indicate to be headquartered in Puerto Rico, foreign to the 50 states.", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"The IRS’s foreign jurisdiction (to Affiant's) and their Roman Civil Law forum in an undisclosed “Constructive Trust” format, is “without” (outside) Affiant's American  National non-statutory domicile and beyond {pronouns[1]} willingness or lawful requirement  with which to comply.", style=indent_list_level_1)
        docx_document.add_paragraph()

        # REASON #22
        paragraph = docx_document.add_paragraph(f"Affiant comes to you herein Commissioner and DIRECTOR in {pronouns[1]} unlimited liability status at peace with you and the IRS and in “good faith,” to cause you to fulfill your duty as public Officers to change Affiant's IRS files and records to a non-taxable status as per {pronouns[1]} enclosed Revocation of Election and AFFIDAVIT in support thereof.", style='List Number 2')
        format_paragraph_keywords(paragraph, ["to change Affiant's IRS files and records to a non-taxable status as per {pronouns[1]} enclosed Revocation of Election and AFFIDAVIT in support thereof."], ["underline"])
        docx_document.add_paragraph()

        # REASON #23
        paragraph = docx_document.add_paragraph(f"Affiant solemnly declares herein that {pronouns[0]} is NOT a “political terrorist” or an “enemy combatant” or a member of any jural society or sovereign political group scheming against {pronouns[1]} country or the Federal Government or the IRS. He has never  participated in any terrorist activities or marches or protests against any Government and {pronouns[0]} has never been involved with any proposed government rebellion, takeover or coups, now or intended in the future.", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Affiant understands that there are certain types of federal citizens and residents with sources of income “within” the U.S. [D.C.] that are “subject to” the IRS's jurisdiction and who are legally obligated to file a Form 1040. Millions of state citizens are liable for filing Form 1040 if they voluntarily “elect” to be taxed as a “federal” person or “U.S. citizen.” Voluntary “servitude” is legal whereby “involuntary servitude” is not.", style=indent_list_level_1)
        docx_document.add_paragraph()

        # REASON #24
        paragraph = docx_document.add_paragraph(f"It is common knowledge among legal researchers such as Judges seated on the  United States Tax Court in D.C., that filing Form 1040 is a voluntary act for State  Citizens of the Union and American Nationals not statutorily connected to the IRS's jurisdiction or the federal government and who are not domiciled in D.C.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #25
        paragraph = docx_document.add_paragraph(f"This is Affiant's notice to you and the IRS that {pronouns[0]} wishes to discontinue volunteering to “donate” {pronouns[1]} personal property through Federal Reserve debt instruments (taxes) to the Treasury effective immediately via {pronouns[1]} Revocation of Election herein.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #26
        paragraph = docx_document.add_paragraph(f"Regardless of whether the IRS is operating under U.S. Bankruptcy laws, the “Trading with the Enemies Act” or the “War Powers Act” (where all federal citizens are legislatively determined to be “enemies of the state” [D.C.]), or the Lieber Code (Martial Law) or D.C. International “municipal” law or by a hidden United Nations Constitution or treaty, or a Papal Bull decree from the Vatican, or an IMF Charter, or the Jesuit General (Illuminati) Zionist, or the P-2 Lodge, or the  Committee of 300 or the Rothschild banking family or ultimately the UPU in Switzerland, any future threats or actions by the IRS to compel Affiant to file Form 1040 will violate too many laws to mention here and especially the District of Columbia's Thirteenth Amendment’s prohibition against involuntary servitude pertaining to slavery. Excessive and illegal taxation is a form of financial slavery.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #27
        paragraph = docx_document.add_paragraph(f"The District of Columbia is a “foreign” corporation with respect to a state of the union [under the original Constitution], 19 Corpus Juris Secondum sec. 883 (2003).", style='List Number 2')
        format_paragraph_keywords(paragraph, ["Corpus Juris Secondum"], ["italic"])
        docx_document.add_paragraph()

        # REASON #28
        paragraph = docx_document.add_paragraph(f"The word “Internal” as in the Internal Revenue Code means “municipal,” i.e., limited to those geographic areas where Congress exercises exclusive legislative authority. In this context, please compare the Federal UCC, which Congress enacted expressly for the District of Columbia. D.C.'s congress and the IRS DO NOT have exclusive jurisdiction over the 50 states of the union.", style='List Number 2')
        format_paragraph_keywords(paragraph, [" Internal ", "i.e."], ["italic"])
        docx_document.add_paragraph()

        # REASON #29
        paragraph = docx_document.add_paragraph(f"In the U.S. Supreme Court case Foley Brothers, Inc v. Filardo, 336 U.S. 281 (1949), the high Court stipulated:", style='List Number 2')
        format_paragraph_keywords(paragraph, ["Foley Brothers, Inc v. Filardo"], ["italic"])
        docx_document.add_paragraph()        
        paragraph = docx_document.add_paragraph(f"“the Cannon of construction which teaches that legislation of Congress, unless a contrary intent appears, is meant to apply only  within the territorial jurisdiction of the United States [meaning D.C.].”", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“the Cannon of construction which teaches that legislation of Congress, unless a contrary intent appears, is meant to apply only  within the territorial jurisdiction of the United States [meaning D.C.].”"], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"The State of Maine’s Supreme Court clarified this issue by explaining our “Right of Election” or “freedom of choice” between two different forms of government –  44 Maine 518 (1859). State Citizens are under no legal or lawful obligation to join  or pledge any allegiance to the foreign legislative democracy [in D.C.].", style=indent_list_level_1)
        docx_document.add_paragraph()

        # REASON #30
        paragraph = docx_document.add_paragraph(f"Because Form 1040 must satisfy IRC 6065, an American National cannot sign  and execute Form 1040 without committing perjury, insofar as Affiant is not a statutory federal citizen or resident alien. A Form 1040 has to be signed by a “federal” citizen “Individual” under 28 U.S.C. 1746 (2). As an American National  and not a “federal” “Individual” I can only autograph documents under 28 U.S.C.  1746 (1), meaning “WITHOUT” (not within) the IRS's and D.C.'s “federal” jurisdiction.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #31
        paragraph = docx_document.add_paragraph(f"Pursuant to many U.S. Supreme Court rulings, the IRS operates under certain Federal Statutes and regulations that are only applicable within the legislative enclave of the District of Columbia; as such, these statutes and regulations comply  with the limited legislative intent of D.C.’s municipal laws. The Sixteenth Amendment was allegedly ratified in response to the high Court’s decision in Pollock v. Farmer’s Loan and Trust Co. In that decision, Congress and the Executive Branch were told by the Supreme Court that they could not “impose” a  “federal” income tax on state Citizens or American Nationals domiciled within the  union (and outside of D.C.'s “federal” jurisdiction) because of well-established Constitutional restrictions against doing so.", style='List Number 2')
        format_paragraph_keywords(paragraph, ["Pollock v. Farmer’s Loan and Trust Co."], ["italic"])
        docx_document.add_paragraph()

        # REASON #32
        paragraph = docx_document.add_paragraph(f"26 U.S.C. 7701(a)(14) defines the term “taxpayer” as any person “subject to” any “internal” [D.C.] revenue tax. For any person to be “subject to” any tax, they must first be under or within the municipal jurisdiction of the federal government, i.e., within that foreign (to the 50 states) 10-mile square area commonly referred to as D.C.", style='List Number 2')
        format_paragraph_keywords(paragraph, ["i.e."], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"As an American National, Affiant does not live or work in D.C. or one of its possessions or territories, thus, the IRS’s jurisdiction does not apply to {pronouns[2]}.", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"A “person” as defined in 26 U.S.C. 7701(a)(1) refers only to “statutory” legal “fictions” “subject to” the federal government and the IRS located in D.C. Affiant  is not a “person” as defined in sec. 7701 (a)(1).", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Affiant does not consent and rebuts herein being defined as a “legal fiction” as “fictions” are dead (on paper only) entities with no inalienable rights and it is a legal impossibility to be both a “living man” presently with inalienable rights while also being considered “dead” by D.C.’s foreign IRS officials at the same time.", style=indent_list_level_1)
        docx_document.add_paragraph()

        # REASON #33
        paragraph = docx_document.add_paragraph(f"As the Supreme Court said in Yick Wo v. Hopkins, 118 U.S. 356 (1886):", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"“Sovereignty itself is, of course, not subject to the law for it is the author and source of the law.”", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“Sovereignty itself is, of course, not subject to the law for it is the author and source of the law.”"], ["italic"])
        docx_document.add_paragraph()

        # REASON #34
        paragraph = docx_document.add_paragraph(f"In the famous case Economy Plumbing & Heating v. U.S., 470 F2d. (1972), this Appellate court declared the existence of two (2) groups related to the Federal income tax. Those groups are “taxpayers” and lawful “non-taxpayers.” Those American Nationals, the lawful “non-taxpayers”, were stated by this Federal Court  to be neither the “subject” nor the “object” of Federal [IRS] revenue laws:", style='List Number 2')
        format_paragraph_keywords(paragraph, ["Economy Plumbing & Heating v. U.S."], ["italic"])        
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"“Revenue laws relate to taxpayers and not to non-taxpayers. No procedures are prescribed for non-taxpayers and no attempt is made to annul any of their Rights or Remedies in due course of law.”", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“Revenue laws relate to taxpayers and not to non-taxpayers. No procedures are prescribed for non-taxpayers and no attempt is made to annul any of their Rights or Remedies in due course of law.”"], ["italic"])
        docx_document.add_paragraph()

        # REASON #35
        paragraph = docx_document.add_paragraph(f"In the U.S. Supreme Court case United States v. Cooper Corporation, 312 U.S.  600 (1941), the court stated:", style='List Number 2')
        format_paragraph_keywords(paragraph, ["Economy Plumbing & Heating v. U.S."], ["italic"])        
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"“Since in common usage the term person does not include the sovereign, statutes not employing the phrase [sovereign] are ordinarily construed to exclude it.”", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“Since in common usage the term person does not include the sovereign, statutes not employing the phrase [sovereign] are ordinarily construed to exclude it.”"], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"The IRC doesn’t mention the word “sovereign” which can only mean that certain “sovereigns” like American Nationals are not within the IRS’s jurisdiction.", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"The IRS’s definitions of “persons” and “taxpayers” make no reference to a tax liability for American Nationals, who are sovereigns by birth in one of the 50 states of the Republic, similar to their parent(s).", style=indent_list_level_1)
        docx_document.add_paragraph()

        # REASON #36
        paragraph = docx_document.add_paragraph(f"To the best of Affiant’s knowledge and understanding, 26 U.S.C. 6013(g)(4)(A) allows “nonresident alien individuals” [similar to an American National] to terminate their previous election to be taxed as a federal “person” or “taxpayer.”", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Affiant purposely uses the term American National to describe {pronouns[1]} political non statutory and non-taxable status affiliation so there is no confusion or mistaken connection between Affiant and IRS terms like “nonresident alien” or “alien” or “nonresident” or “individual” or “resident alien” or “U.S. person” or “U.S. citizen.”", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Affiant declares that {pronouns[0]} is none of these federal government created juristic “statutory terms,” as {pronouns[0]} is first and foremost a God created 'man' living under the laws of God, on the land under  the rule of law, and not living under the law of the  Sea (Admiralty law). By Affiant's right of birth and being a qualified American National, Affiant has no legal obligation to be under the jurisdiction of the IRS when their policies and regulations do not deal with or regulate legal non-taxpayers with no “income” received from a federal government source.", style=indent_list_level_1)
        docx_document.add_paragraph()

        # REASON #37
        paragraph = docx_document.add_paragraph(f"With Affiant's Revocation of Election notice herein - not to be classified as a “taxpayer” from this day forward, {pronouns[0]} rebuts any further “presumptions” by the IRS that {pronouns[0]} is a “taxpayer” and {pronouns[0]} demands said change in {pronouns[1]} tax status to be duly noted and made in all of the IRS's files and databases related to Affiant's tax status.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #38
        paragraph = docx_document.add_paragraph(f"A main purpose of this Notice to you, Commissioner and DIRECTOR, is to give you Affiant's authority and consent herein to recognize everything within this  collective document as Affiant's Revocation of Election and to cause {pronouns[1]} IRS files  to be changed accordingly.", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Affiant endeavors to assist you in your efforts to act on this Revocation of Election notice by including an AFFIDAVIT herein stating facts relating to Affiant's “non-federal” IRS affiliations and {pronouns[1]} “non-taxable” status.", style=indent_list_level_1)
        docx_document.add_paragraph()

        # REASON #39
        paragraph = docx_document.add_paragraph(f"American Nationals and “nonresident alien individuals” (IRS term), with no federal connections and their rights not to be taxed or compelled to make an election to be taxed, was very clearly established by the legislative intent of the Sixteenth Amendment. American Nationals and the people of the 50 states of the union have always been defined as “non-taxpayers” related to Form 1040 (until they volunteered to be taxpayers) and they were explicitly excluded from being “subject to” D.C.'s Sixteenth Amendment and its limited jurisdiction applications.  See legislative “intent” written by President William Taft in the U.S. Senate Congressional records of June 16, 1909 on pages 3344-3345.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #40
        paragraph = docx_document.add_paragraph(f"Affiant has never knowingly performed the functions of “public office,” the statutory definition of a “trade or business” per 26 U.S.C. 7701(a)(21).", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #41
        paragraph = docx_document.add_paragraph(f"As an American National and a non-federal “person,” Affiant has not received  any “taxable” income within the “federal” United States [D.C.] and {pronouns[0]} does not have a “tax year.” The term “income” is not defined anywhere in the IRC according to the Eighth Circuit Court of Appeals. U.S. v. Ballard, 535 F.2nd 400, 404 (8th Circuit, 1976), but this fact is irrelevant to Affiant as nothing written in 26  U.S.C. or the IRC applies to Affiant as a legal non-taxpayer. However the word “income” is defined in the IRC, said definitions only apply to taxpayers and not to  Affiant as a legal non-taxpayer.", style='List Number 2')
        format_paragraph_keywords(paragraph, ["U.S. v. Ballard, 535 F.2nd 400, 404 (8th Circuit, 1976)"], ["italic"])
        docx_document.add_paragraph()

        # REASON #42
        paragraph = docx_document.add_paragraph(f"In IRC section 6013(g) or (h) -“nonresident alien individuals” [and non-taxable American Nationals] “may elect” to volunteer to have their income taxed as a U.S. [D.C.] “resident alien” and thus be obligated to file a Form 1040.", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"The expression “may elect” above is vitally important as it clearly proves that there is no “mandatory” obligation on an American National (state Citizen) to  file a Form 1040 and pay an income tax. A non-taxable American National or “nonresident alien individual” “may elect” to be treated as a “federal” person and  “volunteer” to file a tax return but they have “no mandatory” obligation to do so.", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"When state Citizens and American Nationals (not legally liable to pay income taxes) send their alleged income tax payments to the Treasury with their Form 1040, the Treasury does not call payments received “income taxes.” Instead, the Treasury calls these tax revenues “Donations” to the Treasury so when millions of  people find out they never owed an income tax and demand their (not owed) tax refunds, the U.S. federal government will not be liable to refund their 'not owed' income taxes because “donations” are voluntary gifts and not refundable.", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"When 'non federal' American Nationals, living within the 50 states of the union, mistakenly file their first Form 1040, they have inadvertently and unknowingly “ELECTED” (read volunteered) to be treated as though they were a federal citizen domiciled in D.C. and thus, legally liable to continue to file Form 1040 until said “ELECTION” has been “REVOKED.” See #41 above – non-taxable [American Nationals] and certain other non-taxable Citizens “may elect” (volunteer) to be treated as federal taxpayers (by filing their Form 1040).", style=indent_list_level_1)
        docx_document.add_paragraph()

        # REASON #43
        paragraph = docx_document.add_paragraph(f"The term “United States” is defined in 31 U.S.C. 321(d)(2) as meaning only the federal government in the District of Columbia and not the 50 states of the Union  per 26 U.S.C. 7408(d).", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #44
        paragraph = docx_document.add_paragraph(f"As mentioned in the beginning of this document, the word JURISDICTION and its legal meaning is most important in law, related to who or what is legally obligated to pay Individual Income Taxes (Form 1040). Affiant's years of research unquestionably leads {pronouns[2]} to conclude that; the IRS has only limited “taxing” jurisdiction within the 10-mile square area of D.C. and its territories, possessions,  and federal districts, and the IRC entirely excludes any references to the original Constitutional Republic and the 50 states of the Union and the American Nationals  and natural born state Citizens occupying these union states.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #45
        paragraph = docx_document.add_paragraph(f"IRC 7701(a)(31) Foreign Estate. There are no “Implementing” Regulations  promulgated in the Federal Register imposing any income tax liability upon American Nationals, who like Affiant, work in the private sector and thus, do not derive income from the conduct of a “trade or business,” defined as the performance of the functions of a “public office” working in federal government within the United States [D.C.] per 26 U.S.C. 7408(d) and 7701(a)(39) and who have never made, or subsequently “revoked”, their “election” (to volunteer) to be taxed as a “U.S. person” or as a “U.S. citizen” [meaning D.C.] taxpayer.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #46
        paragraph = docx_document.add_paragraph(f"Affiant is not a “tax protester” as {pronouns[0]} believes in paying all taxes lawfully owed  in {pronouns[1]} jurisdiction. See # 45 above - IRC 7701 is saying that the Estates of state Citizens are considered to be Foreign Estates, meaning “foreign” to the jurisdiction of the IRS - headquartered in that “foreign enclave” being “foreign” to the 50 states of the union and known as the District of Columbia. As a “non-taxable” American National and according to the IRC and your receipt of this ROE, Affiant  cannot be a “tax protester” when there is no income tax imposed upon American Nationals to protest. (See RRA98 re: the IRS or its agents calling people “illegal tax protesters”).", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #47
        paragraph = docx_document.add_paragraph(f"Tax “evasion” is a crime of evading a “lawful” tax. This crime can only be committed by persons who first have a legal liability to pay a tax. Affiant cannot be “evading” income taxes when the IRC clearly defines Affiant (an American National with a ROE sent to the IRS), as a legal “non-taxpayer.”", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #48
        paragraph = docx_document.add_paragraph(f"Affiant is not employed as defined by the Public Salary Tax Act and has no “income” as defined in the IRC - U.S. v. Ballard 535 F, 2d 400, 404 (8th Circuit, 1926).", style='List Number 2')
        format_paragraph_keywords(paragraph, ["U.S. v. Ballard"], ["italic"])
        docx_document.add_paragraph()

        # REASON #49
        paragraph = docx_document.add_paragraph(f"The U.S. Supreme Court said Congress cannot establish a “trade or business” in a state of the union [the Republic] and tax it.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #50
        paragraph = docx_document.add_paragraph(f"Subtitle A income taxes have no effective date of enactment or enforcement published in the Federal Register, a requirement imposed upon the Department of Treasury by 44 U.S.C. 1505 and 26 CFR 601.702(a)(2)(ii), which proves it is not applicable law within the 50 states of the union for American Nationals.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #51
        paragraph = docx_document.add_paragraph(f"The words of the Sixteenth Amendment unequivocally prove that the “federal”  D.C. income tax does not apply to American Nationals not federally connected.", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"“The Congress shall have the power to lay and collect taxes on income, from whatever sources derived, without apportionment among the several States, and without regard to any census or enumeration.”", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“The Congress shall have the power to lay and collect taxes on income, from whatever sources derived, without apportionment among the several States, and without regard to any census or enumeration.”"], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"However, the federal municipal government in D.C. conveniently failed to mention something very important at the end of their Sixteenth Amendment. What the “municipal” D.C. congress knew, or should have known, but failed to mention within “their” private corporation Sixteenth Amendment was:", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"“This power to lay and collect taxes only applies to the National government, officers and employees effectively connected to federal government employment, receiving federal income as wages and those federal citizens effectively domiciled and residing in the District of Columbia, and this Amendment does not apply to state Citizens of the union without (not within) D.C.'s limited taxing jurisdiction and who are not receiving federal government sourced income by being employed by the federal government.”", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“This power to lay and collect taxes only applies to the National government, officers and employees effectively connected to federal government employment, receiving federal income as wages and those federal citizens effectively domiciled and residing in the District of Columbia, and this Amendment does not apply to state Citizens of the union without (not within) D.C.'s limited taxing jurisdiction and who are not receiving federal government sourced income by being employed by the federal government.”"], ["italic"])        
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Adding this simple sentence or clarification to the true intent of D.C.'s Sixteenth Amendment municipal law would have made it crystal clear as to who is and who  isn’t liable for the income tax, as this added sentence reflects the correct legislative intent of the Sixteenth Amendment, well documented in congressional records prior to its enactment and higher Court tax dispute decisions in favor of state Citizens.", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"The word and term “several” states in the Sixteenth Amendment (above) do not mean the 50 states of the union. The “several” states the Sixteenth Amendment is referring to (above) are only D.C. and its “several states” possessions like Guam, American Samoa, and the Virgin Islands etc., also part of the United States [D.C.]  and frequently referred to as the “several” states.", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"The words “without apportionment” (see #51 above) also confirm and prove that the Sixteenth Amendment only refers to the 10-mile square jurisdictional area of D.C. and its territories, because “un-apportioned” income taxes are strictly prohibited in the 50 states of the union under the original Constitution (1787) (as ruled upon in Pollock v. Farmers Loan & Trust Co., Supreme Court decision), but  are legally allowed to be “without apportionment” in the “municipal law” area of D.C. because D.C. is not one of the 50 states of the union, and therefore, not subject to the state Citizen's (of the 50 states) protective restrictions of the original  Constitution (1787 plus Amendments 1789).", style=indent_list_level_1)
        format_paragraph_keywords(paragraph, ["Pollock v. Farmers Loan & Trust Co."], ["italic"])        
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"A “U.S. citizen” and a “U.S. person” are clearly defined today as being “federal” citizens with limited inalienable rights and are the intended targets of the Sixteenth  Amendment - the “federal” (National Government) citizens to be taxable and not the state Citizens of the union. Federal citizens could be (income) taxed because their “federal” employment and income from government employment is considered a “privilege” and privileges can be taxed whereby “inalienable” rights to “earned” property, like money, in the private sector cannot be taxed.", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"As President William Taft said, …”the Sixteenth Amendment will be a tax on the National Government.” There was no question as to whom President Taft was referring – basically, federal government employees only and not state Citizens.", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"The Sixteenth Amendment, even if it was legally ratified, was an Amendment to the District of Columbia's local ten mile area 'municipal' government’s constitution that only applied to the District of Columbia municipal government. D.C.'s constitution was not the same as our country's original Constitution that only had Thirteen Amendments.", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"There was no Sixteenth Amendment in the original Constitution as income taxes on state Citizens [American Nationals] not apportioned were explicitly forbidden in the original Constitution (1787) and they still are forbidden to date. Not all Constitutions are the same. The Sixteenth Amendment did not amend the original  Constitution, it only amended the District of Columbia's private corporation constitution.", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"The D.C. municipal government, being a “foreign” district or enclave to the 50 states of the union, can pass any income tax laws they want, applicable only to government employees, officers, and contractors and federal citizens within the jurisdiction of the District of Columbia, but “their” income tax laws cannot be repugnant to the original Constitution (1787) that forbids un-apportioned income taxes on state Citizens (and American Nationals) living in the 50 states of the union and higher court rulings on tax disputes have confirmed this fact.", style=indent_list_level_1)
        docx_document.add_paragraph()

        # REASON #52
        paragraph = docx_document.add_paragraph(f"Based on Affiant's free will choice to choose {pronouns[1]} political affiliations, take note  that {pronouns[0]} is not now a Fourteenth Amendment “federal” citizen under D.C.’s federal  jurisdiction as this “municipal” Amendment would attempt to define Affiant as a “federal” citizen subject to a “federal” income tax liability.", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"The Fourteenth Amendment (related to the IRS’s jurisdiction) is federal, “non-positive” D.C. local law, foreign law to the union states. The Fourteenth Amendment was enacted to set up a voluntary “Cestui Que Vie” trust relationship  between 'federal' citizens and the Federal government that any state Citizen of the  union states could participate in IF DESIRED and at their option.", style=indent_list_level_1)
        format_paragraph_keywords(paragraph, ["“Cestui Que Vie”"], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"If Affiant accepted D.C.’s municipal law Fourteenth Amendment, it would shift {pronouns[1]}  American National constitutionally protected “rights” authority, into being an indentured Fourteenth Amendment “federal citizen” “within” the federal municipal government's and the IRS's jurisdiction in D.C. Affiant is not now and never intends to be a Fourteenth Amendment citizen.", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"However, members of Congress at the time of the passing of the Fourteenth Amendment were aware of how they were entrapping uninformed state Citizens with D.C's Fourteenth Amendment (choice) to relinquish their “inalienable rights”  for limited government “privileges” to be administered under a municipal government Constructive Trust under Admiralty Law (of the Sea) instead of the Common Law of the “Land” of the union states, which is far more state Citizen-friendly and more closely connected to the original Constitution and the Bill of Rights.", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Therefore, one day before the Fourteenth Amendment was passed, Congress (with  guilty consciences) passed 15 Stat. 249-250, allowing Citizens to remove themselves from the Fourteenth Amendment “public trust jurisdiction” (under the  Doctrine of Parens Patria - the state is the father) if they so desired and that legal  option to not be a Fourteenth Amendment citizen still exists today.", style=indent_list_level_1)
        docx_document.add_paragraph()
        format_paragraph_keywords(paragraph, ["Parens Patria"], ["italic"])

        # REASON #53
        paragraph = docx_document.add_paragraph(f"Affiant thanks the federal government for offering {pronouns[2]} this Fourteenth Amendment option of “cradle to grave” “privileges” if {pronouns[0]} gives them everything {pronouns[0]} owns; however, Affiant respectfully declines the offer. Affiant does not consent to  relinquishing {pronouns[1]} God given inalienable rights with free will agency and dominion,  in exchange for limited or no inalienable rights and limited government privileges  that can be taken away at the whim of government. Government “privileges” can be taken away by governments where God given “inalienable rights” cannot be.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #54
        paragraph = docx_document.add_paragraph(f"In the case of Plessy v. Ferguson 61 U.S. 537 542 (1896), the Supreme Court ruled that:", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"“Slavery implies involuntary servitude – a state of bondage … the control of labor of one man for the benefit of another … the word servitude was intended to prohibit the use of all forms of “involuntary” slavery, of whatever class or name.”", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“Slavery implies involuntary servitude – a state of bondage … the control of labor of one man for the benefit of another … the word servitude was intended to prohibit the use of all forms of “involuntary” slavery, of whatever class or name.”"], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"“The Liberty guaranteed is that of a NATURAL and not of an ARTIFICIAL PERSON” Western Turf Ass'n v. Greenburg, 204 U.S. 359, 27 Sup. Ct. 384, 51 L. Ed. 520.", style=indent_list_level_1)
        format_paragraph_keywords(paragraph, ["“The Liberty guaranteed is that of a NATURAL and not of an ARTIFICIAL PERSON” Western Turf Ass'n v. Greenburg, 204 U.S. 359, 27 Sup. Ct. 384, 51 L. Ed. 520."], ["italic"])
        docx_document.add_paragraph()

        # REASON #55
        paragraph = docx_document.add_paragraph(f"The term “American National” is never used in the IRC because sentient, natural-born living men and women are not statutory / juristic / “U.S. persons” / “U.S. citizens” / “resident aliens” or any other such term as all of these terms are defined by the IRS to be “taxpayers” who either derive income from various sources connected to the federal government in D.C. or who have unwittingly allowed themselves to be “presumed” by the IRS to be voluntary “taxpayers” due  to people filing the IRS's Form 1040 in the past because someone incorrectly told them they had to.", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"American Nationals are not any of the “terms” in #55 above and are not mentioned even once in the IRC. The IRC and the IRS only deal with “taxpayers” and tax law  and regulations related thereto, so there is no reason to mention American Nationals in the IRC. It’s apparent that the IRS recognizes this fact. The IRS has no dealings whatsoever with legal “non-taxpayers” not subject to IRS tax laws and higher court case rulings have proven this point.", style=indent_list_level_1)
        docx_document.add_paragraph()

        # REASON #56
        paragraph = docx_document.add_paragraph(f"A sentient man’s or woman’s (not statutory person's) income tax liability depends mainly on their jurisdiction status, their source of earnings, and their free  will choice to choose their preferred (non-taxable) jurisdiction.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #57
        paragraph = docx_document.add_paragraph(f"In the IRS publication 519:", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"“A non-resident alien [American National] participant who never worked in the U.S. Government in the United States [D.C.]  will not be liable for the U.S. [D.C.] income tax.”", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“A non-resident alien [American National] participant who never worked in the U.S. Government in the United States [D.C.]  will not be liable for the U.S. [D.C.] income tax.”"], ["italic"])
        docx_document.add_paragraph()

        # REASON #58
        paragraph = docx_document.add_paragraph(f"According to 26 U.S.C. 7701(b)(1)(B), a ”nonresident alien” is defined as: one  who is neither a “U.S. citizen” [meaning a D.C. federal citizen] nor a “resident” [of D.C.]. The term “nonresident” means someone who doesn’t “reside” or live in D.C. and the term “alien” basically means someone is “foreign” to D.C.’s jurisdiction if they are a state Citizen of the union under the Constitution (1787), a  free will choice anyone qualified can make.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #59
        paragraph = docx_document.add_paragraph(f"On behalf of a past IRS Commissioner Charles Rossotti, Director Cloonan stated in {pronouns[1]} letter to an American National:", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"“Our system of taxation is dependent on the taxpayer’s belief that the laws they follow apply to everyone….”", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“Our system of taxation is dependent on the taxpayer’s belief that the laws they follow apply to everyone….”"], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Mr. Rossotti, as a past IRS Commissioner, must have known what {pronouns[0]} was talking about when referring to a system based on someone’s “belief” in what the law says, so please be advised that Affiant “believes” that {pronouns[0]}, as an American National, is not subject to or liable for the Subtitle A Form 1040 income tax, or an “excise tax” or a “gift tax” and commands that you honor {pronouns[1]} Revocation of Election by removing Affiant from the IRS’s tax rolls.", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Commissioner Rossotti also stated in a delegated response letter that:", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"“The law itself does not require individuals to file a Form 1040.”", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“The law itself does not require individuals to file a Form 1040.”"], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Affiant appreciates Mr. Rossotti telling the truth on this subject; however, this truth never seems to be taught or trickle down to the lower level IRS agents.", style=indent_list_level_1)
        docx_document.add_paragraph()

        # REASON #60
        paragraph = docx_document.add_paragraph(f"In a letter from Mark L. Forman, Legislative Correspondent, U.S. Senate, dated 6/26/89, Mr. Forman wrote:", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"“Based on the research performed by the Congressional Research Service, there is no provision which specifically and unequivocally requires an individual to pay income taxes.”", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“Based on the research performed by the Congressional Research Service, there is no provision which specifically and unequivocally requires an individual to pay income taxes.”"], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"This does not surprise  me as as our country's founding fathers never intended for state Citizens to be liable for an un-apportioned 'federal' income tax and this kind of tax was strictly prohibited in the original Constitution.", style=indent_list_level_1)
        docx_document.add_paragraph()

        # REASON #61
        paragraph = docx_document.add_paragraph(f"Under oath before Congress in 1953, Dwight E. Avis, Bureau of Internal Revenue stated in part:", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"“Your income tax is a 100% voluntary tax.”", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“Your income tax is a 100% voluntary tax.”"], ["italic"])
        docx_document.add_paragraph()

        # REASON #62
        paragraph = docx_document.add_paragraph(f"Affiant is also not required to file a Form 1040 because Title 26 is not “positive law” as it has never been promulgated in the Federal Register and there are no “implementing regulations” pertaining to a Form 1040-type tax applicable to an American National.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #63
        paragraph = docx_document.add_paragraph(f"When tax regulations are not registered in the Federal Register, this absence means that these “unrecorded” regulations do not apply to state Citizens or American Nationals, as by law, tax “regulations” must be recorded (promulgated) in the Federal Register as a form of “notice” to the public without D.C.’s jurisdiction, or the public (American Nationals and state Citizens) would have no idea of which laws passed by Congress applied to them. Municipal laws passed in D.C. for federal “U.S. citizens” only, do not have to be registered in the Federal Register.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #64
        # THIS IS ONE OF THE REASONS I REMOVED
        paragraph = docx_document.add_paragraph(f"Affiant knows the U.S. Tax Court in D.C. recognizes American Nationals, like  Affiant, as being “without” (not within) the IRS’s jurisdiction. Many American Nationals have dealt with the U.S. Tax Court in D.C. when attacked by the IRS, in  order to get the IRS's (Notice of Deficiency) and certain other IRS claims against them dismissed.", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Affiant knows of a person who hadn’t filed 1040 forms since 2003 when {pronouns[0]} received a Notice of Deficiency (NOD) in July 2014 from the IRS, stating that  {pronouns[0]}  had at least 90 days to open a dispute challenge with the U.S. Tax Court in D.C. regarding this NOD {pronouns[0]} received from the IRS and alleged income taxes due.", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"After sending the U.S. Tax Court a letter explaining {pronouns[1]} American National non-taxable status and not granting the Tax Court any jurisdiction, the Chief Judge Michael B. Thornton issued an ORDER OF DISMISSAL FOR LACK OF JURISDICTION, dated Sept 15, 2014, dismissing all claims the IRS thought they  had against this American National.", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Affiant's friend never had to go to court in D.C. and no attorney was needed to assist {pronouns[2]}. Similar DISMISSALS in this D.C. Tax Court confirm the fact that the IRS recognizes their lack of taxing jurisdiction over American Nationals. This fact  can be confirmed by you by contacting Chief Judge Michael B. Thornton at the U.S. Tax Court in D.C.", style=indent_list_level_1)
        docx_document.add_paragraph()

        # REASON #65
        paragraph = docx_document.add_paragraph(f"Those who work for the IRS are obligated to abide by the U.S. Tax Court decisions – Internal Revenue Manual (IRM) Section 4, 10, 7, 2, 9, 8 (5-14-99) and RRA98.", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Michael L. White, Federal Attorney, Office of the Federal Register, openly stated in {pronouns[1]} legal opinion letter in 1994 that there are no enforcement regulations published in the Federal Register nor is there any published requirement there requiring American Nationals to file or pay an income tax. Mr. White stipulated:", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"“There are no corresponding entries for Title 26.”", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“There are no corresponding entries for Title 26.”"], ["italic"])
        docx_document.add_paragraph()

        # REASON #66
        paragraph = docx_document.add_paragraph(f"28 U.S.C. 7851 (a)(6)(A) states there is no authority for the IRS to use any enforcement action against American Nationals until 26 U.S.C. has been enacted into positive law (meaning it applies to state Citizens or American Nationals only  after it has been promulgated in the Federal Register – making it positive law).", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #67
        paragraph = docx_document.add_paragraph(f"A Statute [related to 26 U.S.C.] is void according to the Supreme Court when it lacks an “implementing” regulation promulgated in the Federal Register and thus cannot be enforced. California Bankers v. Schultz, 416 US 25, 44 39 L. Ed 2nd 912,  94 S. Court.", style='List Number 2')
        format_paragraph_keywords(paragraph, ["California Bankers v. Schultz"], ["italic"])        
        docx_document.add_paragraph()

        # REASON #68
        paragraph = docx_document.add_paragraph(f"In the case United States v. Eaton, the court found it absolutely essential to deal with the corresponding [implementing] regulation. Statutes represent the intent of  the law in general and the regulations related to the Statute more specifically define how the Statute’s intent will be carried out or “implemented.”", style='List Number 2')
        format_paragraph_keywords(paragraph, ["United States v. Eaton"], ["italic"])
        docx_document.add_paragraph()

        # REASON #69
        paragraph = docx_document.add_paragraph(f"In the case of U.S. v. Mersky, 361 US 431, IRC Section 6001 can’t be enforced  without there first being an “implementing” regulation promulgated in the Federal  Register so Section 6001 does not apply to American Nationals.", style='List Number 2')
        format_paragraph_keywords(paragraph, ["U.S. v. Mersky"], ["italic"])
        docx_document.add_paragraph()

        # REASON #70
        paragraph = docx_document.add_paragraph(f"The term “natural person” is not found in the IRC as “natural persons” are flesh and blood living beings and not taxable in general. Statutory law doesn't deal with  “natural persons.” It only deals with millions of man-made “code” violations related to “juristic” fiction entities where there is no actual “harmed party.”", style='List Number 2')
        format_paragraph_keywords(paragraph, ["U.S. v. Mersky"], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Only “fiction” nom-de-guerre entity “constructs” created by the municipal law D.C. government and the IRS can be taxed, assuming “full disclosure” and mutual  agreement is granted, as “privileges” granted from a “fiction” private municipal government and accepted by another “fiction” citizen, may be a binding contract. A living American National man or woman (not fictions) without (not within) D.C.'s  jurisdiction has their God given inalienable “right” to keep their earned or equal “exchanged for labor” private property ($).", style=indent_list_level_1)
        format_paragraph_keywords(paragraph, ["nom-de-guerre"], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"The Oregon Supreme Court was quite clear when it said that the individual, unlike  the corporation, cannot be taxed for the mere privilege of existing, and that the corporation is an artificial entity which owes its existence and charter power to the  State; but the individual's rights to live and own property are natural rights for the  enjoyment of which an excise cannot be imposed. Redfield v. Fisher, 292P 813,819 (1930).", style=indent_list_level_1)
        format_paragraph_keywords(paragraph, ["Redfield v. Fisher"], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"The Tennessee Supreme Court was also quite clear when it said that the right to receive income or earnings is a right belonging to every person, this right cannot be taxed as a privilege. Jack Cole v. MacFarland, 337 S.W. 2D 453, 456 (Tenn. 1960).", style=indent_list_level_1)
        format_paragraph_keywords(paragraph, ["Jack Cole v. MacFarland"], ["italic"])
        docx_document.add_paragraph()

        # REASON #71
        paragraph = docx_document.add_paragraph(f"Affiant declares here that {pronouns[0]} is a natural person man alive and well in the Constitutional Republic of Indiana without (not within) the juristic statutory foreign jurisdiction of the IRS and nothing in the IRC applies to “natural persons”  (unless by election) as said regulations would then be unlawful and repugnant to the Constitution (1787) as President William Taft and many higher court case rulings on income tax issues have so aptly pointed out.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #72
        paragraph = docx_document.add_paragraph(f"26 U.S.C. has never been registered as “positive” law in the Federal Register for the last sixty-one years, thus, 26 U.S.C. is nothing more than “prima facie” (acting as the law until rebutted) “municipal” law and it has no force and effect on  Affiant when rebutted by Affiant and a prior, presumed by the IRS taxable election, has been “revoked” herein by Affiant.", style='List Number 2')
        format_paragraph_keywords(paragraph, ["“prima facie”"], ["italic"])
        docx_document.add_paragraph()

        # REASON #73
        paragraph = docx_document.add_paragraph(f"Affiant is not an “individual” as defined at 5 U.S.C. (a)(2) as this term effectively means a citizen of the United States [D.C.] with a legal domicile on a federal territory and who works for the federal government either directly or indirectly.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #74
        paragraph = docx_document.add_paragraph(f"Affiant is not now nor has {pronouns[0]} ever been involved with the manufacturing of Alcohol, Tobacco or Firearms or any other “trade or business” that might generate  taxable excise income. The Parallel Tables of Authority for IRS enforcement rules  strangely lead directly to the Bureau of Alcohol, Tobacco or Firearms (BATF), a foreign organization with which Affiant has no affiliation and there are no (law) enforcement statutes for the IRS.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #75
        paragraph = docx_document.add_paragraph(f"There are no “regulations” extending to the Commissioner of the IRS or Department of the Treasury, their authority to the 50 union states - 26 CFR 7802(a).", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #76
        paragraph = docx_document.add_paragraph(f"The IRC is only “prima facie” - “color of law” as per 1 USCA 204(a) and only a “presumption” of law if not rebutted. Affiant rebuts herein any claim from the IRS that {pronouns[0]} is a “taxpayer” via {pronouns[1]} Revocation of Election herein.", style='List Number 2')
        format_paragraph_keywords(paragraph, ["“prima facie”"], ["italic"])
        docx_document.add_paragraph()

        # REASON #77
        paragraph = docx_document.add_paragraph(f"The IRS’s definition of the term “taxpayer” means any person subject to any internal revenue tax. By the IRS’s definition, Affiant, as an American National with a Revocation of Election duly sent to you and the IRS, cannot now be classified as  a “taxpayer” as {pronouns[0]} is not subject to any “internal” (to D.C.) revenue tax and you are “duty bound” and required to accept Affiant's Revocation of Election documents and change Affiants's tax records to indicate “non-taxable.”", style='List Number 2')
        format_paragraph_keywords(paragraph, ["“prima facie”"], ["italic"])
        docx_document.add_paragraph()

        # REASON #78
        paragraph = docx_document.add_paragraph(f"Affiant is technically and lawfully a “foreign” estate to the IRS with earnings from sources without the United States [D.C.] and not connected with a “trade or business” within the United States [D.C.] and with no gross income under Subtitle A.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #79
        paragraph = docx_document.add_paragraph(f"In 27 U.S.C. sec. 26.11, the definition of “Revenue Agent” is: “Any duly authorized Commonwealth Internal Revenue Agent” of the Department of the Treasury of Puerto Rico. The Secretary is defined as: Secretary of the Treasury of  Puerto Rico. Affiant does not live in Puerto Rico nor is {pronouns[0]} knowingly connected in  any way to Puerto Rico and Affiant has no recollection of receiving any notices from the IRS related to any contractual connection to Puerto Rico.", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"The IRS appears to be a “trust” domiciled in Puerto Rico as per 31 U.S.C. 1321 (a) (62) and the IRS is not an agency of the federal government as that term is defined  in the Freedom of Information Act (FOIA) and the Administrative Procedures Act in 5 U.S.C. 551(1)(C).", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Affiant fails to understand how the IRS, domiciled in Puerto Rico, can have any “personam,” “venue,” or “subject matter” jurisdiction over {pronouns[2]} without {pronouns[1]} expressed consent, which {pronouns[0]} has never granted to the IRS.", style=indent_list_level_1)
        format_paragraph_keywords(paragraph, ["“personam,”"], ["italic"])
        docx_document.add_paragraph()

        # REASON #80
        paragraph = docx_document.add_paragraph(f"Title 48 U.S.C. plainly states that the entire Internal Revenue Code (IRC) from  start to finish is “generally” made up of “internal revenue laws” which are relevant  to the enforcement of Title III of the National Prohibition Act which only has venue jurisdiction in Puerto Rico and the Virgin Islands.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #81
        paragraph = docx_document.add_paragraph(f"As a sentient man, Affiant's organic name given to {pronouns[2]} at {pronouns[1]} birth is not spelled in all capital letters. A name spelled in all capital letters according to government grammar and writing style manuals means, an “artificial” or “fictitious” or “dead” entity like a corporate person or trustee or partnership, etc.", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"To the extent the IRS claims a right to tax Affiant’s all capital letter nom-de-guerre artificial name, Affiant herein rebuts and denies the IRS's “presumed” right to do so based on never receiving adequate “full disclosure” when the IRS's deceptive and illegal contractual action was “imposed” upon Affiant. All taxing contracts have to include “full disclosure” or they are invalid.", style=indent_list_level_1)
        format_paragraph_keywords(paragraph, ["nom-de-guerre"], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Affiant is aware that the IRS is a corporate “fiction” and that “fictions” can only deal or interact with other “fictions” and not with “natural” living beings and American Nationals without their consent. Affiant does not consent to be converted from a living man to a “fictitious” dead entity (existing on paper only) and {pronouns[0]} does  not consent to being used as a “surety” or “guarantor” for the IRS’s Admiralty Maritime law forum in its actions as a income tax debt collector for “foreign” international bankers when Affiant is a legal “non-taxpayer” with no taxes due.", style=indent_list_level_1)
        docx_document.add_paragraph()

        # REASON #82
        paragraph = docx_document.add_paragraph(f"Because Affiant is not now a “taxpayer” from the moment you received {pronouns[1]} Revocation of Election herein, no IRS regulation related to Subtitle A income taxes pertains to Affiant as an American National. Any references to IRS regulations made by Afiant herein are merely being made for the purpose of substantiating Affiant's non-taxable status.", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Affiant's using said IRC references herein is in no way meant by Affiant to be interpreted or construed by you or the IRS to mean that Affiant grants or acknowledges IRS's jurisdiction over {pronouns[2]} just because {pronouns[0]} is referencing IRS regulations to explain {pronouns[1]} non-taxable status.", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Not one sentence or paragraph in the IRC or 26 U.S.C. has any legal application to Affiant as an American National with no income from a government source and Affiant's quoting IRC references is only for explaining {pronouns[1]} “non-taxable” status and not {pronouns[1]} agreement that the IRC or 26 U.S.C. is the final law authority related to {pronouns[1]} Revocation of Election and {pronouns[1]} non-taxable status.", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"For example: IRC sec. 1.1-1 entitled:", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Income tax on individuals....“imposes an income tax on the income of every individual who is a citizen or resident of the United States [D.C.].", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["Income tax on individuals....“imposes an income tax on the income of every individual who is a citizen or resident of the United States [D.C.]."], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Affiant thanks the IRS for making this clear statement as to “who is” and “who is  not” liable for the tax. Affiant, as an American National, is not an “individual” based on definitions in various U.S.C. Titles which thankfully confirm {pronouns[1]} adamant position of not being an “individual” “subject to” or “liable for” the Form 1040 “Individual” federal income tax as {pronouns[0]} is neither an “individual” (meaning statutory  “federal” person) nor a “citizen” or a “resident” of the United States [D.C] as {pronouns[0]} does not “reside” in D.C. or any federal area or zone related thereto. The word “income” is not defined anywhere in the IRC, thus, IRC 1.1-1 is void on its face for vagueness.", style=indent_list_level_1)
        docx_document.add_paragraph()

        # REASON #83
        paragraph = docx_document.add_paragraph(f"Living and breathing men and women and non-Fourteenth Amendment American Nationals are not “individuals” as defined in the U.S. Codes. Title 26 U.S.C. 7701 (a) (31) basically says that an American National's Estate is a tax-exempt foreign estate or trust.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #84
        paragraph = docx_document.add_paragraph(f"According to 5 U.S.C. (a), under the Administrative Procedures Act, an “individual” means a “U.S. citizen,” a fictitious entity [adhesioned to D.C.’s federal jurisdiction] with no inalienable rights, i.e., a juristic statutory “person” effectively domiciled in the District of Columbia, regardless of whether said “individual” actually lives in D.C. Affiant is not an “individual” and {pronouns[0]} rebuts all attempted references by the IRS to being an “individual.” 5 U.S.C., therefore, does  not apply to Affiant as a American National man.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #85
        paragraph = docx_document.add_paragraph(f"The Buck Act has converted most state Citizens into being “domiciled” within  D.C. without their awareness of this fact. Affiant does not live in a Buck Act “federal zone” represented by a two capital letter federal zone designation for each of the 50 States, meaning, Affiant does not reside or live in the “federal zone” designated by the two capital letters of {soujourn_state_abbreviations} for the States of {soujourn_state_names}.", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Instead, Affiant's domicile is “Non-domestic” (not within a D.C “federal zone”) in the {soujourn_republic_names}, zip code exempt, under the protections of the Constitution (1787) and Affiant's God given inalienable rights to choose {pronouns[1]} jurisdiction.", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Affiant's inalienable rights do not need a Constitution or any other legal document  source to secure, protect, or determine {pronouns[1]} American National status because sovereignty is the source of law from God and needs no document to confirm it. Yick Wo v. Hopkins, 118 U.S. 356 (1886)", style=indent_list_level_1)
        format_paragraph_keywords(paragraph, ["Yick Wo v. Hopkins"], ["italic"])
        docx_document.add_paragraph()

        # REASON #86
        paragraph = docx_document.add_paragraph(f"High level IRS personnel regularly state that the Federal individual income tax  is “voluntary” for certain non-taxable people. Because Affiant is not a “taxpayer” as per {pronouns[1]} Revocation of Election herein, {pronouns[0]} agrees with the IRS’s position that Form 1040 income tax filings are “voluntary” and herein gives you notice of {pronouns[1]} intention to not volunteer to file a Form 1040 in the future.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #87
        paragraph = docx_document.add_paragraph(f"In Flora v. United States, 362 U.S. says in part:", style='List Number 2')
        format_paragraph_keywords(paragraph, ["Flora v. United States"], ["italic"])        
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"“...regarding congressional intent, our system of taxation is based upon voluntary assessment and payment not  upon distraint.”", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“...regarding congressional intent, our system of taxation is based upon voluntary assessment and payment not  upon distraint.”"], ["italic"])        
        docx_document.add_paragraph()

        # REASON #88
        paragraph = docx_document.add_paragraph(f"In Long v. Rasmussen, revenue laws relate to “taxpayers” not to “non-taxpayers.” Delima v. Bidwell, 182 U.S. 176, 179 and Gerth v. United States, 132 F. Supp. 894 (1955) ruled similarly.", style='List Number 2')
        format_paragraph_keywords(paragraph, ["Long v. Rasmussen", "Delima v. Bidwell", "Gerth v. United States"], ["italic"])
        docx_document.add_paragraph()

        # REASON #89
        paragraph = docx_document.add_paragraph(f"As a living man with unlimited liability and not requesting any government “privileges” or “benefits,” Affiant does not have or own a social security card nor does {pronouns[0]} have a social security number.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #90
        paragraph = docx_document.add_paragraph(f"“Congress [IRS] might have territorial or legislative jurisdiction solely over U.S. [“federal” citizens] “residing” within Washington, D.C., the federal enclaves inside the states and outside the continental area of the United States.” Berman v.  Parker, 343 U.S. 26, 75 S. Ct. 98 (1954); and Cincinnati Soap Co. v. United States, 301 U.S. 303, 57 S.Ct. 764 (1937).", style='List Number 2')
        format_paragraph_keywords(paragraph, ["“Congress [IRS] might have territorial or legislative jurisdiction solely over U.S. [“federal” citizens] “residing” within Washington, D.C., the federal enclaves inside the states and outside the continental area of the United States.” Berman v.  Parker","Cincinnati Soap Co. v. United States"], ["italic"])
        docx_document.add_paragraph()

        # REASON #91
        paragraph = docx_document.add_paragraph(f"On current U.S. Passport Applications, it says:", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"“U.S. Passports either in Book  or Card Format, are issued only to U.S. Citizens or Non-Citizen Nationals.”", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“U.S. Passports either in Book  or Card Format, are issued only to U.S. Citizens or Non-Citizen Nationals.”"], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"A Non-Citizen National can be defined as a state Citizen or an American National and a “U.S. Citizen” means a D.C. “federal” citizen as there are basically only two types of citizens.", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"This is further evidence and conclusive proof that the federal government's Passport Office in D.C. recognizes and acknowledges the clear distinction between a “federal” citizen under D.C.’s (and possibly IRS’s) jurisdiction and a state Citizen “National” (or an American National) without (not within) D.C.’s federal (taxing) jurisdiction.", style=indent_list_level_1)
        docx_document.add_paragraph()

        # REASON #92
        paragraph = docx_document.add_paragraph(f"“The sovereign is merely sovereign by {pronouns[1]} very existence. The rule in America is that the American people are the sovereigns.” Kemper v. State 138 Southwest 1025 (1911), pg. 1043, sec. 33.", style='List Number 2')
        format_paragraph_keywords(paragraph, ["“The sovereign is merely sovereign by {pronouns[1]} very existence. The rule in America is that the American people are the sovereigns.” Kemper v. State"], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"If a person is an American National with God given inalienable rights to life, liberty, and the pursuit of happiness, then they cannot also be a “federal” citizen who accepts benefits and privileges from a federal government as the law prohibits dual “elections.”", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Any natural-born state Citizen or American National not connected to the “foreign” federal government in D.C. and its “municipal” law (meaning limited and local) has the inalienable right to choose their political and taxing jurisdiction. Affiant's American National status has no known or intended connection to the IRS’s “federal” and “foreign” jurisdiction.", style=indent_list_level_1)
        docx_document.add_paragraph()

        # REASON #93
        paragraph = docx_document.add_paragraph(f"Affiant's “non-taxable” status is explained by the following U.S. Supreme Court case:", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"“The individual may stand upon {pronouns[1]} constitutional [secured] Rights as  a [non federal] citizen. He is entitled to carry on {pronouns[1]} private business in {pronouns[1]} own way. His power to contract is unlimited. He owes no such duty [to submit {pronouns[1]} books and records for examination] to the State [IRS and federal government] since {pronouns[0]} receives nothing therefrom beyond the protection of {pronouns[1]} life and property. His Rights are such as existed by the law of the land [common law] long antecedent to  the organization of the State [IRS], and can only be taken from {pronouns[2]} by due process  of law and in accordance with the Constitution [1787]. Among {pronouns[1]} Rights are a refusal to incriminate himself and the immunity of himself and {pronouns[1]} property from arrest or seizure except under a warrant of the law. He owes nothing to the public  so long as {pronouns[0]} does not trespass upon their Rights.” Hale v. Henkel, 201 U.S. 43 at  47 (1905).", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“The individual may stand upon {pronouns[1]} constitutional [secured] Rights as  a [non federal] citizen. He is entitled to carry on {pronouns[1]} private business in {pronouns[1]} own way. His power to contract is unlimited. He owes no such duty [to submit {pronouns[1]} books and records for examination] to the State [IRS and federal government] since {pronouns[0]} receives nothing therefrom beyond the protection of {pronouns[1]} life and property. His Rights are such as existed by the law of the land [common law] long antecedent to  the organization of the State [IRS], and can only be taken from {pronouns[2]} by due process  of law and in accordance with the Constitution [1787]. Among {pronouns[1]} Rights are a refusal to incriminate himself and the immunity of himself and {pronouns[1]} property from arrest or seizure except under a warrant of the law. He owes nothing to the public  so long as {pronouns[0]} does not trespass upon their Rights.” Hale v. Henkel"], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Since 1905, Hale v. Henkel has been cited within the federal and state appellate court systems over a thousand times and none of the issues of law in this case have  ever been overturned.", style=indent_list_level_1)
        format_paragraph_keywords(paragraph, ["Hale v. Henkel"], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Affiant, as an American National, has never asked for “permission” from the federal government and IRS, nor has {pronouns[0]} required permission from the government,  hence, {pronouns[0]} has no “duty” to the federal government. As an American National with  God given inalienable rights, protected by the Constitution (1787), Affiant does not have a liability to file a Form 1040 as the IRC and 26 U.S.C. specifically and expressly does not include or define Affiant to be a “taxpayer.”", style=indent_list_level_1)
        docx_document.add_paragraph()

        # REASON #94
        paragraph = docx_document.add_paragraph(f"Under the IRS “Restructuring and Reformation Act of 1998” (RRA98), the  burden of proof that someone is an “illegal tax protester” falls completely on the IRS and the IRS, under RRA98, is prohibited from calling someone an “illegal tax  protester” in one's IRS records and any IRS agent who does so can be terminated from {pronouns[1]} job (fired) for this unproven or false allegation or for any other Internal Revenue Manual violation. Affiant does not protest any “lawful” taxes owed in {pronouns[1]}  jurisdiction.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #95
        paragraph = docx_document.add_paragraph(f"The U.S. Supreme Court ruled, a law applicable in Washington, D.C. was not applicable in San Antonio [Republic of Texas] because it did not conform to Constitutional restrictions. U.S. v. Lopez 115 S. ct. 1624 (1995).", style='List Number 2')
        format_paragraph_keywords(paragraph, ["U.S. v. Lopez"], ["italic"])
        docx_document.add_paragraph()

        # REASON #96
        paragraph = docx_document.add_paragraph(f"The “federal” income tax laws that are applicable to federal citizens in Washington D.C. are not applicable to “non-federal” state Citizens or American Nationals that do not conform to the original Constitution restrictions and the Bill  of Rights.", style='List Number 2')
        format_paragraph_keywords(paragraph, ["U.S. v. Lopez"], ["italic"])
        docx_document.add_paragraph()

        # REASON #97
        paragraph = docx_document.add_paragraph(f"The above stated Internal Revenue Code laws, rules, regulations, and court case rulings confirm and substantiate Affiant's Revocation of Election and the IRS's internal records must now show John Q. Public's “non-taxable” status.", style='List Number 2')
        docx_document.add_paragraph()

        # REASON #98
        paragraph = docx_document.add_paragraph(f"“It is not the function of our Government to keep the citizen from falling into error, it is the function of the Citizen to keep the Government from falling into error.” American Communications Associations v. Douds, 339 U.S. 382,442, (1950)", style='List Number 2')
        format_paragraph_keywords(paragraph, ["“It is not the function of our Government to keep the citizen from falling into error, it is the function of the Citizen to keep the Government from falling into error.” American Communications Associations v. Douds"], ["italic"])
        docx_document.add_paragraph()

        # REASON #99
        paragraph = docx_document.add_paragraph(f"“The IRS is not a government agency. It is an agency of the International Monetary Fund (IMF).” (Diversified Metals Products v. IRS et al. CV-93-405E U.S.D.C.D. I., Public Law 94-564, Senate Report 94-1148 pg. 5967, Reorg. Plan no. 26, Public Law 102-391.", style='List Number 2')
        format_paragraph_keywords(paragraph, ["“The IRS is not a government agency. It is an agency of the International Monetary Fund (IMF).” (Diversified Metals Products v. IRS et al. CV-93-405E U.S.D.C.D. I., Public Law 94-564, Senate Report 94-1148 pg. 5967, Reorg. Plan no. 26, Public Law 102-391."], ["italic"])
        docx_document.add_paragraph()

        # REASON #100
        paragraph = docx_document.add_paragraph(f"Title 26 IS NOT LAW. The Internal Revenue Code defines the contract between the IRS and the individual. 26 USC 7806(b) says that Title 26 is not law, as we read:", style='List Number 2')
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"“No inference, implication or presumption of legislative construction [meaning Law] shall be drawn or made by reason of the location or grouping of any particular section or provision or portion of this title.” [emphasis added].", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“No inference, implication or presumption of legislative construction [meaning Law] shall be drawn or made by reason of the location or grouping of any particular section or provision or portion of this title.”"], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"26 USC sec. 6331 is not imposed upon American Nationals.", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"“Levy and distraint (a) Authority of Secretary ... Levy may be made upon the accrued salary or wages  of any officer, employee, or elected official, of the United States, the District of Columbia, or any agency or instrumentality of the United States or the District of  Columbia....”", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“Levy and distraint (a) Authority of Secretary ... Levy may be made upon the accrued salary or wages  of any officer, employee, or elected official, of the United States, the District of Columbia, or any agency or instrumentality of the United States or the District of  Columbia....”"], ["italic"])
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"Affiant cannot be levied against as Affiant believes {pronouns[0]} is not any of  the above persons and furthermore, nothing in 26 U.S.C. applies to non-taxpayers.", style=indent_list_level_1)
        docx_document.add_paragraph()
        paragraph = docx_document.add_paragraph(f"“The U.S. Supreme Court affirms that no one can “elect” [choose] to be treated as a resident of a particular place [Federal location] for the purpose of taxation without also having a factual presence [living] in that location.” Texas v. Florida,  306 U.S. 398 (1939).", style=indent_list_level_2)
        format_paragraph_keywords(paragraph, ["“The U.S. Supreme Court affirms that no one can “elect” [choose] to be treated as a resident of a particular place [Federal location] for the purpose of taxation without also having a factual presence [living] in that location.” Texas v. Florida"], ["italic"])
        docx_document.add_paragraph()
        #########################################################################################################
        #               END OF LETTER OF SUPPLEMENTAL INFORMATION
        #########################################################################################################

    def create_w_8ben_pdf(self, w_8ben_pdf_file_path, first_given_name, middle_given_name, family_name, man_or_woman, selected_states, street_address, city, zip_code, mailing_state, social_security_number, irs_commissioner, local_irs_service_center_street_address, local_irs_service_center_city, local_irs_service_center_state_list, local_irs_service_center_zip, sojourn_states_list, mailing_address_state_list, republic_of_birth_list, notary_state_list, notary_county, country_of_citizenship, date_of_birth, include_ohio_state_edits):
        #########################################################################################################
        #                       START OF DEFINING VARIABLES
        #########################################################################################################
        # Extract the "State Name" from each dictionary
        soujourn_state_names_list = [state_dict["State Name"] for state_dict in sojourn_states_list]
        soujourn_state_names = ", ".join(soujourn_state_names_list)

        # Extract the "State Abbreviations" from each dictionary
        soujourn_state_abbreviations_list = [state_dict["State Abbreviation"] for state_dict in sojourn_states_list]
        soujourn_state_abbreviations = ", ".join(soujourn_state_abbreviations_list)

        # Extract the "State " from each dictionary
        soujourn_republic_name_list = [state_dict["Republic Name"] for state_dict in sojourn_states_list]
        soujourn_republic_names = ", ".join(soujourn_republic_name_list)

        mailing_address_republic_name = mailing_address_state_list[0]['Republic Name']
        mailing_address_state_name = mailing_address_state_list[0]['State Name']
        mailing_address_state_abbreviation = mailing_address_state_list[0]['State Abbreviation']

        republic_of_birth_name = republic_of_birth_list[0]['Republic Name']
        republic_of_birth_state = republic_of_birth_list[0]['State Name']
        republic_of_birth_state_abbreviation = republic_of_birth_list[0]['State Abbreviation']

        local_irs_service_center_republic_name = local_irs_service_center_state_list[0]['Republic Name']
        local_irs_service_center_state__name = local_irs_service_center_state_list[0]['State Name']
        local_irs_service_center_state_abbreviation = local_irs_service_center_state_list[0]['State Abbreviation']

        notary_republic_name = notary_state_list[0]['Republic Name']
        notary_state_name = notary_state_list[0]['State Name']
        notary_state_abbrivation = notary_state_list[0]['State Abbreviation']

        vessel_name = first_given_name + ' ' + middle_given_name + ' ' + family_name
        live_name = first_given_name + '-' + middle_given_name + ': ' + family_name

        # Create variables for titlecased and uppercase versions of the name
        titlecased_name = vessel_name.title()
        uppercase_name = vessel_name.upper()

        # Determine the pronouns based on the 'are_you' variable
        if man_or_woman == 'Man':
            pronouns = ("he", "his", "him")
            gender = man_or_woman.lower()
        elif man_or_woman == 'Woman':
            pronouns = ("she", "her", "her")
            gender = man_or_woman.lower()

        # Get the current date in the desired format
        current_date = datetime.now().strftime("%m-%d-%Y")
        #########################################################################################################
        #                       END OF DEFINING VARIABLES
        #########################################################################################################

        # Get the path to the directory containing your executable
        app_dir = os.path.dirname(os.path.realpath(__file__))

        if include_ohio_state_edits == 'Yes':
            # Input PDF file path
            pdf_input_path = os.path.join(app_dir, 'assets', 'fw8ben_with_ohio_edits.pdf')
            # Access the form fields on the page
            field_values = {
                'f_1[0]': f'{live_name}, beneficiary',                  # Part 1, Field 1: Name of individual who is the beneficial owner 
                'f_2[0]': f'{country_of_citizenship}',                  # Part 1, Field 2: Country of citizenship
                'f_3[0]': 'rural free delivery',                        # Part 1, Field 3: Permanent residence address (street, apt. or suite no., or rural route). Do not use a P.O. box or in-care-of address. 
                'f_4[0]': f'{city}, {mailing_address_state_name}',      # Part 1, Field 3: City or town, state or province. Include postal code where appropriate.
                'f_5[0]': f'{republic_of_birth_state}',                 # Part 1, Field 3: Country
                'f_6[0]': f'c/o rr {street_address}',                   # Part 1, Field 4: Mailing address (if different from above)
                'f_7[0]': f'{city}, {mailing_address_state_name}',      # Part 1, Field 4: City or town, state or province. Include postal code where appropriate.
                'f_8[0]': f'{republic_of_birth_state}',                 # Part 1, Field 4: Country
                'f_9[0]': f'{social_security_number}',                  # Part 1, Field 5: U.S. taxpayer identification number (SSN or ITIN), if required (see instructions)
                'f_10[0]': '',                                          # Part 1, Field 6a: Foreign tax identifying number (see instructions)
                'c1_01[0]': '/0',                                       # Part 1, Field 6b  (checkbox): Check if FTIN not legally required
                'f_11[0]': '',                                          # Part 1, Field 7: Reference number(s) (see instructions)
                'f_12[0]': f'{date_of_birth}',                          # Part 1, Field 8: Date of birth (MM-DD-YYYY) (see instructions)
                'f_13[0]': 'N/A',                                       # Part 2, Field 9: I certify that the beneficial owner is a resident of...
                'f_14[0]': 'N/A',                                       # Part 2, Field 10: Special rates and conditions / Article and paragraph
                'f_15[0]': 'N/A',                                       # Part 2, Field 10: Special rates and conditions / % rate
                'f_16[0]': 'N/A',                                       # Part 2, Field 10: Special rates and conditions / (specify type of income):
                'f_17[0]': '',                                          # Part 2, Field 10: Explain the additional conditions in the Article and paragraph the beneficial owner meets to be eligible for the rate of withholding (this is the short line right after the word "withholding:")
                'f_18[0]': '\t\t\tN/A',                                 # Part 2, Field 10: Explain the additional conditions in the Article and paragraph the beneficial owner meets to be eligible for the rate of withholding (this is the longer line below)
                'c1_02[0]': '/1',                                       # Part 3, Checkbox: I certify that I have the capacity to sign for the person identified on line 1 of this form.
                'f_20[0]': '',                                          # Part 3, Signature Line
                'Date[0]': f'{current_date}',                           # Part 3, Date: Date (MM-DD-YYYY)
                'f_21[0]': '',                                          # Part 3, Print name of signer
                'Text1': 'Non-Resident Alien',                          # Part 1, Identification of "Non-Resident Alien"
                'Text2': 'Non-Resident Alien',                          # Part 1, Name of individual who is the "Non-Resident Alien"
                'Text3': '  domicile',                                    # Part 1, Permanent "domicile" address
                'Text4': 'Speical\n\nPrivate\n\nand\n\nPriority',                # Margin text, "Speical Private and Priority"
                'Text5': 'non-resident alien owner',                    # Part 2, Field 9: "non-resident alien owner"
                'Text6': 'non-resident alien owner',                    # Part 3, "non-resident alient owner"
                'Text7': '© All Right Reserved 28 U.S.C. § 1746(1)',    # Signature line "© All Right Reserved 28 U.S.C. § 1746(1)"
                'Text8': 'non-resident alient owner',                   # Below signature line "non-resident alient owner"
                'Text9': 'non-resident alient owner',                   # Below signature line (or individual authorized to sign for "non-resident alient owner")
                'Text10': 'In Exclusive Equity',                        # Print name of signer
            }
        else:
            # Input PDF file path
            pdf_input_path = os.path.join(app_dir, 'assets', 'fw8ben.pdf')
            # Access the form fields on the page
            field_values = {
                'f_1[0]': f'{live_name}, beneficiary',                  # Part 1, Field 1: Name of individual who is the beneficial owner 
                'f_2[0]': f'{country_of_citizenship}',                  # Part 1, Field 2: Country of citizenship
                'f_3[0]': 'rural free delivery',                        # Part 1, Field 3: Permanent residence address (street, apt. or suite no., or rural route). Do not use a P.O. box or in-care-of address. 
                'f_4[0]': f'{city}, {mailing_address_state_name}',      # Part 1, Field 3: City or town, state or province. Include postal code where appropriate.
                'f_5[0]': f'{republic_of_birth_state}',                 # Part 1, Field 3: Country
                'f_6[0]': f'c/o rr {street_address}',                   # Part 1, Field 4: Mailing address (if different from above)
                'f_7[0]': f'{city}, {mailing_address_state_name}',      # Part 1, Field 4: City or town, state or province. Include postal code where appropriate.
                'f_8[0]': f'{republic_of_birth_state}',                 # Part 1, Field 4: Country
                'f_9[0]': f'{social_security_number}',                  # Part 1, Field 5: U.S. taxpayer identification number (SSN or ITIN), if required (see instructions)
                'f_10[0]': '',                                          # Part 1, Field 6a: Foreign tax identifying number (see instructions)
                'c1_01[0]': '/0',                                       # Part 1, Field 6b  (checkbox): Check if FTIN not legally required
                'f_11[0]': '',                                          # Part 1, Field 7: Reference number(s) (see instructions)
                'f_12[0]': f'{date_of_birth}',                          # Part 1, Field 8: Date of birth (MM-DD-YYYY) (see instructions)
                'f_13[0]': 'N/A',                                       # Part 2, Field 9: I certify that the beneficial owner is a resident of...
                'f_14[0]': 'N/A',                                       # Part 2, Field 10: Special rates and conditions / Article and paragraph
                'f_15[0]': 'N/A',                                       # Part 2, Field 10: Special rates and conditions / % rate
                'f_16[0]': 'N/A',                                       # Part 2, Field 10: Special rates and conditions / (specify type of income):
                'f_17[0]': '',                                          # Part 2, Field 10: Explain the additional conditions in the Article and paragraph the beneficial owner meets to be eligible for the rate of withholding (this is the short line right after the word "withholding:")
                'f_18[0]': '\t\t\tN/A',                                 # Part 2, Field 10: Explain the additional conditions in the Article and paragraph the beneficial owner meets to be eligible for the rate of withholding (this is the longer line below)
                'c1_02[0]': '/1',                                       # Part 3, Checkbox: I certify that I have the capacity to sign for the person identified on line 1 of this form.
                'f_20[0]': '',                                          # Part 3, Signature Line
                'Date[0]': f'{current_date}',                           # Part 3, Date: Date (MM-DD-YYYY)
                'f_21[0]': '',                                          # Part 3, Print name of signer
            }


        # Create a PdfFileWriter to write the updated PDF
        pdf_writer = PyPDF2.PdfWriter()

        # List to store form field names
        field_names = []

        # Open the input PDF file in read-binary mode
        with open(pdf_input_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            # Check if the PDF has form fields
            if pdf_reader.is_encrypted:
                pdf_reader.decrypt("")  # If the PDF is encrypted, provide a password here

            # Extract form field names
            if '/AcroForm' in pdf_reader.trailer:
                acro_form = pdf_reader.trailer['/AcroForm']
                for field in acro_form['/Fields']:
                    field_names.append(field['/T'])

            # Now you have a list of form field names
            print("Form Field Names:", field_names)

            # Access the first page of the PDF (modify the page index as needed)
            page = pdf_reader.pages[0]


            # Update the form field values
            pdf_writer.update_page_form_field_values(page, field_values)
            pdf_writer.add_page(page)

        # Save the updated PDF to the output file
        with open(w_8ben_pdf_file_path, 'wb') as output_pdf:
            pdf_writer.write(output_pdf)


    #########################################################################################################
    #                       END OF create_supporting_evidence FUNCTION
    #########################################################################################################


def main():
    app = QApplication(sys.argv)
    ex = PDFGeneratorApp()
    sys.exit(app.exec_())

if __name__ == '__main__':
    main()